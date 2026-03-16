#!/usr/bin/env python3
"""
HMS Enterprise — Product Document (Enterprise Grade)
Rich-themed DOCX generator with strong branding.
Author: Ashish Kumar Satyam | Organization: TechDigital WishTree
Date: 16 March 2026
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime
import os

# ─── Output ───
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, "docs", "enterprise")
os.makedirs(OUTPUT_DIR, exist_ok=True)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "HMS_Enterprise_Product_Document.docx")

# ─── Company Branding ───
COMPANY_FULL = "TechDigital WishTree"
COMPANY_SHORT = "TGWT"
PRODUCT_NAME = "HMS Enterprise"
DOC_ID = "TGWT-HMS-PRD-2026-001"
AUTHOR = "Ashish Kumar Satyam"
DOC_DATE = "16 March 2026"
VERSION = "1.0"

# ─── Brand Colors ───
NAVY       = RGBColor(27, 58, 92)
DARK_NAVY  = RGBColor(15, 38, 64)
TEAL       = RGBColor(13, 115, 119)
TEAL_LIGHT = RGBColor(16, 185, 129)
GOLD       = RGBColor(196, 154, 42)
AI_BLUE    = RGBColor(79, 70, 229)
AI_PURPLE  = RGBColor(124, 58, 237)
WHITE      = RGBColor(255, 255, 255)
BLACK      = RGBColor(26, 32, 44)
GRAY       = RGBColor(113, 128, 150)
LIGHT_GRAY = RGBColor(237, 242, 247)
RED        = RGBColor(229, 62, 62)
ORANGE     = RGBColor(237, 137, 54)
GREEN      = RGBColor(56, 161, 105)

# Hex versions
H_NAVY      = "1B3A5C"
H_DARK_NAVY = "0F2640"
H_TEAL      = "0D7377"
H_GOLD      = "C49A2A"
H_AI_BLUE   = "4F46E5"
H_AI_PURPLE = "7C3AED"
H_LIGHT     = "F7FAFC"
H_LIGHTER   = "EDF2F7"
H_WHITE     = "FFFFFF"
H_BLUE_TINT = "EBF8FF"
H_PURPLE_TINT = "FAF5FF"
H_GREEN_TINT = "F0FFF4"
H_GOLD_TINT  = "FFFFF0"
H_RED_TINT   = "FFF5F5"
H_ORANGE_TINT = "FFFAF0"
H_TEAL_TINT  = "E6FFFA"

# ═══════════════════════════════════════════════════════════════════
#  THEMING HELPERS
# ═══════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="D4D4D4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), sz)
            el.set(qn('w:color'), color)
            borders.append(el)
    tcPr.append(borders)

def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        margins.append(el)
    tcPr.append(margins)

def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_paragraph_borders(paragraph, bottom_color="1B3A5C", bottom_sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), bottom_sz)
    bottom.set(qn('w:color'), bottom_color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_borders_left(paragraph, color_hex, sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), color_hex)
    left.set(qn('w:space'), '8')
    pBdr.append(left)
    pPr.append(pBdr)

def set_paragraph_borders_full(paragraph, color_hex, sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for edge in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color_hex)
        el.set(qn('w:space'), '4')
        pBdr.append(el)
    pPr.append(pBdr)

def add_run(para, text, size=11, bold=False, italic=False, color=BLACK, font_name='Calibri', light=False):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    fn = 'Calibri Light' if light else font_name
    run.font.name = fn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), fn)
    rFonts.set(qn('w:hAnsi'), fn)
    rFonts.set(qn('w:cs'), fn)
    return run

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

# ═══════════════════════════════════════════════════════════════════
#  THEMED COMPONENTS
# ═══════════════════════════════════════════════════════════════════

def add_heading_h1(doc, text, icon=""):
    """Navy full-width bar + gold accent strip."""
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(28)
    bar.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(bar, H_DARK_NAVY)
    add_run(bar, "  ", size=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_shading(p, H_DARK_NAVY)
    p.paragraph_format.left_indent = Cm(0.5)
    if icon:
        add_run(p, icon + "  ", size=20, color=GOLD)
    add_run(p, text.upper(), size=20, bold=True, color=WHITE)

    accent = doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after = Pt(14)
    set_paragraph_shading(accent, H_GOLD)
    add_run(accent, " ", size=3)
    return p

def add_heading_h2(doc, text, icon=""):
    """Left teal border + light blue background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_borders_left(p, H_TEAL, "20")
    set_paragraph_shading(p, H_BLUE_TINT)
    p.paragraph_format.left_indent = Cm(0.4)
    if icon:
        add_run(p, icon + "  ", size=14, color=TEAL)
    add_run(p, text, size=14, bold=True, color=NAVY)
    return p

def add_heading_h3(doc, text, icon=""):
    """Gold underline + diamond bullet."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)
    set_paragraph_borders(p, H_GOLD, "4")
    if icon:
        add_run(p, icon + "  ", size=11, color=GOLD)
    else:
        add_run(p, "\u25C6  ", size=10, color=GOLD)
    add_run(p, text, size=12, bold=True, color=NAVY)
    return p

def add_heading_h4(doc, text):
    """Small teal-accented subheading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "\u25B8  ", size=9, color=TEAL)
    add_run(p, text, size=11, bold=True, color=DARK_NAVY)
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.45
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    add_run(p, text, size=10.5, color=RGBColor(45, 55, 72), light=True)
    return p

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(1.2 + indent_level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    marker = "\u2022  "
    add_run(p, marker, size=10.5, color=TEAL)
    if bold_prefix:
        add_run(p, bold_prefix + " ", size=10.5, bold=True, color=NAVY)
    add_run(p, text, size=10.5, color=RGBColor(45, 55, 72), light=True)
    return p

def add_numbered(doc, number, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    add_run(p, f"{number}. ", size=10.5, bold=True, color=TEAL)
    if bold_prefix:
        add_run(p, bold_prefix + " \u2014 ", size=10.5, bold=True, color=NAVY)
    add_run(p, text, size=10.5, color=RGBColor(45, 55, 72), light=True)
    return p

def add_callout(doc, text, bg_hex=H_BLUE_TINT, icon="\u2139\uFE0F", border_color=H_TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    set_paragraph_shading(p, bg_hex)
    set_paragraph_borders_left(p, border_color, "18")
    p.paragraph_format.line_spacing = 1.45
    if icon:
        add_run(p, icon + "  ", size=12)
    add_run(p, text, size=10.5, italic=True, color=RGBColor(45, 55, 72), light=True)
    return p

def add_key_value_box(doc, label, value, bg_hex=H_LIGHT):
    """Inline key-value styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    set_paragraph_shading(p, bg_hex)
    add_run(p, f"  {label}: ", size=10, bold=True, color=NAVY)
    add_run(p, value, size=10, color=RGBColor(45, 55, 72), light=True)
    return p

def add_section_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "\u2500" * 50, size=8, color=LIGHT_GRAY)

def add_stat_row(doc, stats):
    """Row of stat boxes: (value, unit, label)."""
    t = doc.add_table(rows=2, cols=len(stats))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, unit, label) in enumerate(stats):
        c_top = t.rows[0].cells[i]
        c_top.text = ""
        set_cell_shading(c_top, H_DARK_NAVY)
        set_cell_margins(c_top, top=80, bottom=20)
        p = c_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, val, size=24, bold=True, color=RGBColor(167, 139, 250))
        if unit:
            add_run(p, f" {unit}", size=10, color=RGBColor(200, 200, 200))
        c_bot = t.rows[1].cells[i]
        c_bot.text = ""
        set_cell_shading(c_bot, H_NAVY)
        set_cell_margins(c_bot, top=20, bottom=60)
        p2 = c_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p2, label.upper(), size=7, bold=True, color=RGBColor(180, 190, 210))
    doc.add_paragraph()
    return t

def add_feature_card(doc, title, description, status="Day 1", phase="Phase 1"):
    """Small feature card with status badge."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.6)
    set_paragraph_shading(p, H_LIGHT)
    set_paragraph_borders_left(p, H_TEAL, "14")
    add_run(p, f"  {title}", size=10.5, bold=True, color=NAVY)
    status_color = GREEN if status == "Day 1" else AI_PURPLE
    add_run(p, f"  [{status}]", size=8.5, bold=True, color=status_color)
    if description:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Cm(1.0)
        add_run(p2, description, size=9.5, color=GRAY, light=True)
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg=H_DARK_NAVY, stripe=True):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        add_run(p, h, size=9, bold=True, color=WHITE)
        set_cell_shading(cell, header_bg)
        set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
        for edge in ['top', 'bottom', 'left', 'right']:
            set_cell_borders(cell, **{edge: 'single'}, color=header_bg)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_cell_margins(cell, top=30, bottom=30, left=60, right=60)
            if val in ("Yes", "Full", "Day 1", "Core", "Included"):
                add_run(p, val, size=9, bold=True, color=GREEN)
            elif val in ("No", "None", "N/A", "Not Available"):
                add_run(p, val, size=9, bold=True, color=RED)
            elif val in ("Partial", "Basic", "Limited", "Optional"):
                add_run(p, val, size=9, bold=True, color=ORANGE)
            elif "Phase" in str(val):
                add_run(p, val, size=9, bold=True, color=AI_PURPLE)
            else:
                add_run(p, val, size=9, color=BLACK, light=True)
            if stripe and r_idx % 2 == 1:
                set_cell_shading(cell, H_LIGHT)
            set_cell_borders(cell, bottom='single', color="E2E8F0", sz="2")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_page_header(doc):
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, DOC_ID, size=7, color=GRAY)
    add_run(hp, f"  |  ", size=7, color=LIGHT_GRAY)
    add_run(hp, COMPANY_SHORT, size=7, bold=True, color=TEAL)

def add_page_footer(doc):
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    set_paragraph_borders(fp, H_LIGHTER, "2")
    add_run(fp, COMPANY_FULL, size=7, bold=True, color=NAVY)
    add_run(fp, f"  |  {PRODUCT_NAME}  |  ", size=7, color=GRAY)
    add_run(fp, "CONFIDENTIAL", size=7, bold=True, color=GOLD)

# ═══════════════════════════════════════════════════════════════════
#  DOCUMENT SECTIONS
# ═══════════════════════════════════════════════════════════════════

import docx.enum.text

def build_cover_page(doc):
    """Enterprise cover page with strong branding."""
    # Top navy band
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_shading(p, H_DARK_NAVY)
        add_run(p, " ", size=6)

    # Company identity
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p, H_DARK_NAVY)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, COMPANY_FULL.upper(), size=14, bold=True, color=GOLD, font_name='Calibri')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p2, H_DARK_NAVY)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, "Technology  \u2022  Digital Transformation  \u2022  Healthcare Innovation", size=9, color=RGBColor(160, 174, 192))

    # Gold divider
    g = doc.add_paragraph()
    g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(g, H_GOLD)
    add_run(g, " ", size=4)

    # Title block
    for _ in range(2):
        sp = doc.add_paragraph()
        set_paragraph_shading(sp, H_DARK_NAVY)
        add_run(sp, " ", size=14)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(title_p, H_DARK_NAVY)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    add_run(title_p, "HMS ENTERPRISE", size=36, bold=True, color=WHITE)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(sub_p, H_DARK_NAVY)
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(8)
    add_run(sub_p, "Product Document", size=22, color=RGBColor(167, 139, 250))

    desc_p = doc.add_paragraph()
    desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(desc_p, H_DARK_NAVY)
    desc_p.paragraph_format.space_before = Pt(4)
    desc_p.paragraph_format.space_after = Pt(4)
    add_run(desc_p, "Cloud-Native Multi-Tenant Hospital Management Platform", size=12, color=RGBColor(160, 174, 192), light=True)

    desc2 = doc.add_paragraph()
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(desc2, H_DARK_NAVY)
    desc2.paragraph_format.space_before = Pt(0)
    desc2.paragraph_format.space_after = Pt(4)
    add_run(desc2, "12 Integrated Modules  \u2022  AI/ML Innovation  \u2022  ABDM & DPDP Compliant", size=10, color=RGBColor(130, 145, 165), light=True)

    for _ in range(2):
        sp = doc.add_paragraph()
        set_paragraph_shading(sp, H_DARK_NAVY)
        add_run(sp, " ", size=14)

    # Teal accent bar
    ta = doc.add_paragraph()
    set_paragraph_shading(ta, H_TEAL)
    add_run(ta, " ", size=3)

    # Metadata table
    meta_t = doc.add_table(rows=5, cols=2)
    meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document ID", DOC_ID),
        ("Author", AUTHOR),
        ("Organization", COMPANY_FULL),
        ("Date", DOC_DATE),
        ("Version", VERSION),
    ]
    for i, (label, value) in enumerate(meta_data):
        c0 = meta_t.rows[i].cells[0]
        c1 = meta_t.rows[i].cells[1]
        c0.text = ""
        c1.text = ""
        set_cell_shading(c0, H_LIGHTER)
        set_cell_margins(c0, top=40, bottom=40, left=100, right=60)
        set_cell_margins(c1, top=40, bottom=40, left=60, right=100)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p0, label.upper(), size=8, bold=True, color=NAVY)
        p1 = c1.paragraphs[0]
        add_run(p1, value, size=9, color=BLACK, light=True)
        set_cell_borders(c0, bottom='single', color="E2E8F0", sz="2")
        set_cell_borders(c1, bottom='single', color="E2E8F0", sz="2")
    meta_t.columns[0].width = Cm(5)
    meta_t.columns[1].width = Cm(9)

    # Classification band
    doc.add_paragraph()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(cp, H_GOLD_TINT)
    set_paragraph_borders_full(cp, H_GOLD, "4")
    add_run(cp, "  CONFIDENTIAL  \u2014  ", size=9, bold=True, color=GOLD)
    add_run(cp, f"Property of {COMPANY_FULL}. Unauthorized distribution prohibited.", size=8, color=GRAY, light=True)

    doc.add_page_break()


def build_toc(doc):
    """Table of Contents page."""
    add_heading_h1(doc, "Table of Contents", "\u2263")

    toc_items = [
        ("01", "Executive Summary"),
        ("02", "Product Vision & Strategy"),
        ("03", "Platform Architecture"),
        ("04", "Module Specifications (12 Modules)"),
        ("05", "RBAC & Security Framework"),
        ("06", "API Design & Integration Layer"),
        ("07", "AI/ML Innovation Roadmap"),
        ("08", "Infrastructure & DevOps"),
        ("09", "Regulatory Compliance"),
        ("10", "Go-to-Market Strategy"),
        ("11", "Revenue & Pricing Model"),
        ("12", "Implementation Roadmap"),
        ("13", "Appendices & References"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1.0)
        set_paragraph_borders(p, H_LIGHTER, "2")
        add_run(p, num, size=12, bold=True, color=TEAL)
        add_run(p, f"    {title}", size=11, color=NAVY)

    doc.add_page_break()


def build_executive_summary(doc):
    add_heading_h1(doc, "Executive Summary", "\u2726")

    add_stat_row(doc, [
        ("12", "", "Integrated Modules"),
        ("$8.5B", "", "India Market 2030"),
        ("70+", "", "API Endpoints"),
        ("99.9%", "", "Uptime SLA"),
    ])

    add_body(doc,
        f"{PRODUCT_NAME} is a cloud-native, multi-tenant hospital management platform engineered by "
        f"{COMPANY_FULL} to transform healthcare delivery across India's rapidly digitizing hospital ecosystem. "
        f"Built on a modern microservices architecture with Spring Boot, Apache Kafka, PostgreSQL, "
        f"and deployed on AWS EKS, the platform delivers 12 fully integrated modules spanning clinical operations, "
        f"pharmacy, laboratory, finance, HR, inventory, business intelligence, patient engagement, "
        f"audit/compliance, IPD management, scheduling, and platform administration."
    )

    add_callout(doc,
        "HMS Enterprise is positioned as the first Indian-origin SaaS platform to offer ambient clinical AI, "
        "real-time clinical decision support (CDSS), and AI-powered revenue cycle management \u2014 "
        "addressing 18 validated market gaps that no incumbent currently fills.",
        H_PURPLE_TINT, "\u26A1", H_AI_PURPLE
    )

    add_heading_h2(doc, "Strategic Differentiators", "\u2605")
    differentiators = [
        ("Schema-per-Tenant Isolation:", "PostgreSQL schema-level data isolation with Row-Level Security \u2014 zero data leakage between tenants."),
        ("Event-Driven Architecture:", "Apache Kafka backbone processing 10,000+ events/second with Socket.IO real-time UI updates."),
        ("ABDM-Native:", "Health ID (ABHA) integration, HL7 FHIR R4 interoperability, and bidirectional health record exchange from Day 1."),
        ("AI-First Roadmap:", "12 AI/ML modules across 4 phases \u2014 from ambient documentation to agentic clinical workflows."),
        ("Regulatory Compliance:", "DPDP Act 2023, NABH 6th Edition (Jan 2025), and ABDM 2026 mandatory guidelines built into the core."),
    ]
    for bold_p, text in differentiators:
        add_bullet(doc, text, bold_p)

    add_section_divider(doc)
    doc.add_page_break()


def build_product_vision(doc):
    add_heading_h1(doc, "Product Vision & Strategy", "\u2B50")

    add_heading_h2(doc, "Vision Statement")
    add_callout(doc,
        "To become India's most trusted cloud-native hospital management platform \u2014 empowering "
        "every hospital, from 50-bed nursing homes to 2000+ bed multi-specialty chains, with "
        "intelligent, interoperable, and regulation-ready digital infrastructure.",
        H_GOLD_TINT, "\u2728", H_GOLD
    )

    add_heading_h2(doc, "Mission")
    add_body(doc,
        "Deliver an enterprise-grade, modular SaaS platform that eliminates operational inefficiency, "
        "automates clinical documentation, enables data-driven decision-making, and ensures seamless "
        "regulatory compliance \u2014 while maintaining the highest standards of data security and patient privacy."
    )

    add_heading_h2(doc, "Target Segments")
    segments = [
        ("Tier 1 \u2014 Multi-Specialty Chains:", "500+ bed hospitals, 10+ branches, enterprise tier ($15\u201325K/month). Full module deployment with dedicated AI features."),
        ("Tier 2 \u2014 Mid-Size Hospitals:", "100\u2013500 beds, professional tier ($5\u201315K/month). Core modules with optional AI add-ons."),
        ("Tier 3 \u2014 Small Hospitals & Clinics:", "20\u2013100 beds, starter tier ($999\u20135K/month). Essential modules with growth path."),
        ("Tier 4 \u2014 Government & Public Health:", "District/PHC hospitals, custom tier. ABDM-compliant, vernacular UI, offline-first capability."),
    ]
    for bold_p, text in segments:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "Product Principles")
    principles = [
        ("Security First:", "Zero-trust architecture, field-level encryption for PII, MFA enforcement for all admin roles."),
        ("Modularity:", "Any module can be activated/deactivated per tenant without affecting others."),
        ("Interoperability:", "FHIR R4 native, ABDM-ready, HL7v2 backward compatibility."),
        ("Performance:", "Sub-500ms API P95 latency, 10K+ concurrent users per tenant."),
        ("Accessibility:", "WCAG 2.1 AA compliant, PWA for mobile-first India, vernacular language support."),
    ]
    for bold_p, text in principles:
        add_bullet(doc, text, bold_p)

    add_section_divider(doc)
    doc.add_page_break()


def build_architecture(doc):
    add_heading_h1(doc, "Platform Architecture", "\u2699")

    add_heading_h2(doc, "Architecture Overview")
    add_body(doc,
        "HMS Enterprise follows a cloud-native microservices architecture pattern with clear separation "
        "of concerns across presentation, API gateway, business logic, data persistence, and infrastructure layers. "
        "Each module operates as an independent service with its own bounded context, communicating "
        "via Apache Kafka for asynchronous events and REST/gRPC for synchronous operations."
    )

    # Architecture stack table
    make_table(doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Presentation", "React 18 + TypeScript + PWA", "Responsive UI with offline capability"],
            ["API Gateway", "Spring Cloud Gateway + Kong", "Rate limiting, auth, routing, load balancing"],
            ["Business Logic", "Spring Boot 3.x (Java 21)", "12 microservices with domain-driven design"],
            ["Messaging", "Apache Kafka 3.x", "Event streaming, async processing, audit trail"],
            ["Cache", "Redis 7.x Cluster", "Session, RBAC cache, real-time data"],
            ["Search", "Elasticsearch 8.x", "Full-text search, BI analytics, log aggregation"],
            ["Database", "PostgreSQL 16 (Schema-per-tenant)", "ACID-compliant, JSONB for clinical data"],
            ["File Storage", "AWS S3 + CloudFront", "Medical images, reports, documents"],
            ["Orchestration", "AWS EKS (Kubernetes)", "Container orchestration, auto-scaling"],
            ["CI/CD", "GitHub Actions + ArgoCD", "GitOps pipeline, 8-stage deployment"],
            ["Monitoring", "Prometheus + Grafana + ELK", "Observability, alerting, log aggregation"],
            ["AI/ML", "Python FastAPI + Claude API", "Clinical AI services with PII de-identification"],
        ],
        col_widths=[3.5, 5.5, 7.0]
    )

    add_heading_h2(doc, "Multi-Tenancy Model")
    add_callout(doc,
        "3-Layer Defense: Schema-level isolation (PostgreSQL) + Row-Level Security policies + "
        "Application-layer tenant_id injection. Subdomain-based routing (tenant.hmssaas.in) with "
        "Flyway-managed schema migrations per tenant.",
        H_TEAL_TINT, "\u26E8", H_TEAL
    )

    add_heading_h3(doc, "Tenant Isolation Guarantees")
    guarantees = [
        ("Schema Isolation:", "Each tenant gets a dedicated PostgreSQL schema \u2014 no shared tables, no cross-tenant queries."),
        ("RLS Policies:", "Database-level Row-Level Security as a second line of defense, even if application logic fails."),
        ("Network Isolation:", "Kubernetes namespace-level network policies prevent cross-service data leakage."),
        ("Encryption:", "AES-256 at rest, TLS 1.3 in transit, tenant-specific encryption keys for PII fields."),
    ]
    for bold_p, text in guarantees:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "Event-Driven Architecture")
    make_table(doc,
        ["Topic Namespace", "Example Events", "Consumers"],
        [
            ["hms.clinical", "consultation.created, prescription.signed, vitals.recorded", "Pharmacy, Lab, Billing, Audit"],
            ["hms.pharmacy", "dispensation.completed, stock.low, drug.interaction.alert", "Clinical, Inventory, Finance"],
            ["hms.lab", "sample.collected, result.validated, report.generated", "Clinical, Patient Portal, Billing"],
            ["hms.finance", "invoice.generated, payment.received, tpa.claim.submitted", "Patient Portal, BI, Audit"],
            ["hms.hr", "attendance.marked, leave.approved, payroll.processed", "Finance, BI, Admin"],
            ["hms.audit", "record.accessed, consent.granted, data.exported", "Compliance, Admin, Security"],
        ],
        col_widths=[3.0, 7.0, 5.0]
    )

    add_section_divider(doc)
    doc.add_page_break()


def build_modules(doc):
    add_heading_h1(doc, "Module Specifications", "\u2B22")

    add_body(doc,
        "HMS Enterprise comprises 12 fully integrated modules, each designed as an independent microservice "
        "with well-defined APIs, event contracts, and data ownership boundaries. Below is the detailed "
        "specification for each module."
    )

    modules = [
        {
            "num": "01",
            "name": "OPD & Clinical Operations",
            "icon": "\U0001F3E5",
            "desc": "Complete outpatient workflow from registration to consultation, including SOAP-based EMR, token queue management, e-prescriptions, and clinical notes.",
            "features": [
                ("Patient Registration & Search:", "Demographics, ABHA linking, document upload, duplicate detection."),
                ("Token Queue Management:", "Auto-generation, priority override, real-time display, estimated wait time."),
                ("SOAP-Based EMR:", "Chief complaint, history, examination, assessment, plan \u2014 with clinical templates."),
                ("E-Prescriptions:", "Drug search (CIMS), interaction checking, dosage calculator, print/digital delivery."),
                ("Vitals Recording:", "BP, pulse, SpO2, temperature, BMI \u2014 with trend charts and alert thresholds."),
                ("Referral Management:", "Inter-department, external referrals, follow-up scheduling."),
            ]
        },
        {
            "num": "02",
            "name": "Pharmacy Management",
            "icon": "\U0001F48A",
            "desc": "End-to-end pharmacy operations with FIFO batch tracking, drug interaction alerts, GST-compliant billing, and controlled substance management.",
            "features": [
                ("Prescription Fulfillment:", "Auto-receive from EMR, substitution approval, partial dispensing."),
                ("Inventory Tracking:", "Batch-level FIFO, expiry alerts (30/60/90 days), auto-reorder triggers."),
                ("Drug Interaction Engine:", "Real-time alerts during prescription and dispensing."),
                ("GST-Compliant Billing:", "HSN codes, slab-wise taxation, credit note management."),
                ("Controlled Substance Log:", "Schedule H/H1 tracking, register maintenance, audit trail."),
            ]
        },
        {
            "num": "03",
            "name": "Laboratory Information System",
            "icon": "\U0001F52C",
            "desc": "NABL-aligned lab workflows with barcode tracking, dual validation, auto-analyzer integration, and patient portal result delivery.",
            "features": [
                ("Order Management:", "Doctor-initiated, walk-in, package orders with priority flags."),
                ("Sample Lifecycle:", "Collection \u2192 barcode label \u2192 transport \u2192 processing \u2192 validation \u2192 report."),
                ("Dual Validation:", "Technician entry + pathologist verification before release."),
                ("Auto-Analyzer Integration:", "HL7/ASTM bidirectional interface for 50+ analyzer models."),
                ("Reference Ranges:", "Age/gender-specific with critical value flagging and auto-notification."),
            ]
        },
        {
            "num": "04",
            "name": "Finance & Billing",
            "icon": "\U0001F4B0",
            "desc": "Multi-module bill aggregation with TPA/insurance workflows, GST compliance, payment gateway integration, and real-time revenue dashboards.",
            "features": [
                ("Bill Aggregation:", "Auto-consolidate charges from OPD, pharmacy, lab, IPD, and procedures."),
                ("TPA & Insurance:", "Pre-auth workflow, claim submission, follow-up tracking, settlement reconciliation."),
                ("Payment Gateway:", "Razorpay/PhonePe integration, UPI, card, net banking, EMI options."),
                ("GST Compliance:", "GSTR-1/3B auto-filing data, HSN mapping, e-invoicing, credit notes."),
                ("Revenue Dashboard:", "Daily/weekly/monthly collections, outstanding analysis, department-wise P&L."),
            ]
        },
        {
            "num": "05",
            "name": "HR & Payroll",
            "icon": "\U0001F465",
            "desc": "Complete workforce management with biometric attendance, shift scheduling, PF/ESIC deductions, and performance tracking.",
            "features": [
                ("Employee Lifecycle:", "Onboarding \u2192 attendance \u2192 leave \u2192 appraisal \u2192 separation."),
                ("Biometric Attendance:", "Integration with ZKTeco/Realtime devices, geo-fencing for mobile punch."),
                ("Shift Management:", "Rotating shifts, swap requests, overtime calculation, duty rosters."),
                ("Payroll Engine:", "CTC structuring, PF/ESIC/PT/TDS deductions, bank file generation."),
                ("Leave Management:", "Policy-based accrual, encashment, compensatory off, holiday calendar."),
            ]
        },
        {
            "num": "06",
            "name": "Inventory & Procurement",
            "icon": "\U0001F4E6",
            "desc": "Consumable and equipment management with auto-reorder, vendor management, purchase order workflows, and GRN processing.",
            "features": [
                ("Inventory Tracking:", "Item-level, batch-level, department-wise stock with barcode support."),
                ("Auto-Reorder:", "Min/max levels, ABC analysis, lead-time based reorder point calculation."),
                ("Purchase Orders:", "Multi-level approval, vendor comparison, rate contracts."),
                ("GRN Processing:", "Goods receipt, quality check, invoice matching, return management."),
                ("Asset Management:", "Equipment register, AMC tracking, depreciation, maintenance schedules."),
            ]
        },
        {
            "num": "07",
            "name": "Business Intelligence",
            "icon": "\U0001F4CA",
            "desc": "Elasticsearch-powered dashboards with 50+ KPIs, department analytics, trend analysis, and executive reporting.",
            "features": [
                ("Executive Dashboard:", "Revenue, footfall, bed occupancy, surgery count, lab turnaround time."),
                ("Department Analytics:", "OPD flow, pharmacy sales, lab TAT, IPD ALOS, HR attendance rates."),
                ("Trend Analysis:", "YoY/MoM comparisons, seasonal patterns, forecasting models."),
                ("Custom Reports:", "Drag-and-drop report builder, scheduled delivery, export (PDF/Excel)."),
                ("Benchmark Comparison:", "NABH metrics, industry averages, peer hospital benchmarking."),
            ]
        },
        {
            "num": "08",
            "name": "Patient Portal (PWA)",
            "icon": "\U0001F4F1",
            "desc": "Progressive web app for patient self-service \u2014 appointment booking, report access, prescription history, and digital payments.",
            "features": [
                ("Self-Service Booking:", "Doctor search, slot selection, rescheduling, cancellation."),
                ("Health Records:", "Lab reports, prescriptions, discharge summaries, vaccination records."),
                ("Digital Payments:", "Bill view, online payment, receipt download, payment history."),
                ("Notifications:", "Appointment reminders, lab results ready, bill generated alerts."),
                ("Feedback System:", "Post-visit surveys, NPS scoring, complaint registration."),
            ]
        },
        {
            "num": "09",
            "name": "Audit & Compliance",
            "icon": "\U0001F512",
            "desc": "Immutable audit logging, DPDP Act consent management, NABH compliance tracking, and 7-year record retention.",
            "features": [
                ("Immutable Audit Trail:", "Every data access, modification, and export logged with timestamp and user context."),
                ("Consent Management:", "DPDP Act 2023 compliant consent collection, storage, and withdrawal workflows."),
                ("NABH Compliance:", "6th Edition (Jan 2025) indicator tracking, gap analysis, evidence management."),
                ("Data Retention:", "7-year policy enforcement, automated archival, right-to-erasure workflow."),
                ("Security Monitoring:", "Failed login tracking, suspicious access patterns, auto-lockout policies."),
            ]
        },
        {
            "num": "10",
            "name": "IPD Management",
            "icon": "\U0001F6CF",
            "desc": "Inpatient lifecycle from admission to discharge \u2014 bed management, nursing records, MAR, clinical pathways, and discharge planning.",
            "features": [
                ("Bed Management:", "Real-time bed board, ward/room/bed hierarchy, housekeeping status integration."),
                ("Admission Workflow:", "Emergency/planned admission, deposit collection, bed allocation."),
                ("Nursing Records:", "Nursing assessment, care plans, I/O charting, fall risk scoring."),
                ("MAR (Medication Administration):", "Scheduled medication tracking, nurse verification, missed-dose alerts."),
                ("Discharge Planning:", "Checklist-based discharge, summary generation, follow-up scheduling, billing clearance."),
            ]
        },
        {
            "num": "11",
            "name": "Appointment Scheduling",
            "icon": "\U0001F4C5",
            "desc": "Real-time slot management with multi-channel booking, no-show tracking, and intelligent scheduling optimization.",
            "features": [
                ("Slot Management:", "Doctor-wise, department-wise availability with buffer time configuration."),
                ("Multi-Channel Booking:", "Web portal, mobile PWA, WhatsApp bot, front-desk, IVR."),
                ("Queue Optimization:", "Dynamic token assignment, estimated wait time, priority escalation."),
                ("No-Show Analytics:", "Pattern detection, overbooking algorithms, reminder automation."),
                ("Resource Scheduling:", "OT booking, procedure rooms, diagnostic equipment slot management."),
            ]
        },
        {
            "num": "12",
            "name": "Platform Administration",
            "icon": "\u2699\uFE0F",
            "desc": "Tenant lifecycle management, schema provisioning, cross-tenant analytics, feature flags, and system health monitoring.",
            "features": [
                ("Tenant Onboarding:", "Automated schema creation, seed data, admin user provisioning, branding setup."),
                ("Feature Flags:", "Per-tenant module activation, A/B testing, gradual rollout control."),
                ("Cross-Tenant Analytics:", "Aggregate metrics (anonymized), platform health, usage patterns."),
                ("System Health:", "Service status, database metrics, Kafka lag, API latency monitoring."),
                ("Billing & Subscription:", "Plan management, usage metering, invoice generation, payment tracking."),
            ]
        },
    ]

    for mod in modules:
        add_heading_h2(doc, f"Module {mod['num']}: {mod['name']}", mod.get('icon', ''))
        add_body(doc, mod['desc'])
        for bold_p, text in mod['features']:
            add_bullet(doc, text, bold_p)
        doc.add_paragraph()  # spacer

    add_section_divider(doc)
    doc.add_page_break()


def build_rbac_security(doc):
    add_heading_h1(doc, "RBAC & Security Framework", "\U0001F6E1")

    add_heading_h2(doc, "6-Tier Role Hierarchy")
    make_table(doc,
        ["Level", "Role", "Scope", "MFA Required", "Session TTL"],
        [
            ["L0", "SuperAdmin", "Platform-wide (all tenants)", "Yes", "15 min"],
            ["L1", "TenantAdmin", "Single tenant (all branches)", "Yes", "30 min"],
            ["L2", "Branch Lead", "Single branch (all departments)", "Yes", "30 min"],
            ["L3", "Department Head", "Single department", "Optional", "60 min"],
            ["L4", "Staff", "Assigned resources only", "Optional", "60 min"],
            ["L5", "Patient", "Own records only", "No", "8 hours"],
        ],
        col_widths=[1.5, 3.0, 4.5, 2.5, 2.5]
    )

    add_heading_h2(doc, "Permission Model")
    add_body(doc,
        "200+ granular permissions organized as Resource \u00D7 Action \u00D7 Field (e.g., prescription:UPDATE, "
        "patient:READ:medicalHistory). Permission resolution follows a 3-step cascade: collect effective "
        "permissions from role hierarchy, evaluate the request against permissions, and apply field-level "
        "masking for READ operations on sensitive data."
    )

    add_heading_h2(doc, "Authentication & Authorization")
    auth_items = [
        ("JWT Bearer Tokens:", "15-minute access tokens + 8-hour refresh via HttpOnly secure cookies."),
        ("MFA Enforcement:", "Mandatory for L0\u2013L2 roles; OTP via SMS/WhatsApp; TOTP app support."),
        ("IdP Integration:", "Keycloak-based identity provider with SAML/OIDC federation support."),
        ("Session Management:", "Redis-backed session store with role-based TTLs and pub/sub invalidation."),
        ("API Rate Limiting:", "Per-tenant limits \u2014 500 req/min (Starter), 2000 req/min (Professional), 5000 req/min (Enterprise)."),
    ]
    for bold_p, text in auth_items:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "Data Protection")
    add_callout(doc,
        "Field-level encryption (AES-256-GCM) for all PII fields: Aadhaar, mobile, email, medical records. "
        "Encryption keys are tenant-specific and rotated every 90 days. All AI/ML API calls use PII "
        "de-identification \u2014 patient data is stripped before external API calls and restored post-processing.",
        H_RED_TINT, "\U0001F510", "E53E3E"
    )

    add_section_divider(doc)
    doc.add_page_break()


def build_api_design(doc):
    add_heading_h1(doc, "API Design & Integration Layer", "\U0001F517")

    add_heading_h2(doc, "API Standards")
    add_key_value_box(doc, "Base URL", "https://api.hmspro.in/v1")
    add_key_value_box(doc, "Tenant-Scoped", "https://api.hmspro.in/v1/tenants/{tenant_id}")
    add_key_value_box(doc, "Specification", "OpenAPI 3.1 (Swagger)")
    add_key_value_box(doc, "Response Envelope", "{success, data, meta, errors}")
    add_key_value_box(doc, "Pagination", "Cursor-based with configurable page_size (default 20, max 100)")
    add_key_value_box(doc, "Versioning", "URI-based (/v1, /v2) with 12-month deprecation policy")

    add_heading_h2(doc, "Key API Endpoints (70+)")
    make_table(doc,
        ["Module", "Endpoint", "Method", "Auth Level"],
        [
            ["OPD", "/patients", "GET/POST", "L3+"],
            ["OPD", "/consultations/{id}", "GET/PUT", "L3+"],
            ["OPD", "/prescriptions", "POST", "L3+"],
            ["Pharmacy", "/dispensations", "POST", "L4+"],
            ["Pharmacy", "/stock/alerts", "GET", "L3+"],
            ["Lab", "/lab-orders", "POST", "L3+"],
            ["Lab", "/lab-results/{id}/validate", "PUT", "L3+"],
            ["Finance", "/invoices", "GET/POST", "L3+"],
            ["Finance", "/payments", "POST", "L4+"],
            ["IPD", "/admissions", "POST", "L3+"],
            ["IPD", "/beds/availability", "GET", "L4+"],
            ["Scheduling", "/appointments", "GET/POST", "L5+"],
            ["Admin", "/tenants/{id}/config", "PUT", "L1"],
            ["AI", "/ai/discharge-summary", "POST", "L3+"],
        ],
        col_widths=[2.5, 5.0, 2.5, 2.0]
    )

    add_heading_h2(doc, "Integration Capabilities")
    integrations = [
        ("ABDM/FHIR:", "HL7 FHIR R4 resources, ABHA health ID linking, bidirectional health record exchange."),
        ("Payment Gateways:", "Razorpay, PhonePe, Paytm \u2014 webhook-based payment confirmation."),
        ("SMS/WhatsApp:", "Twilio/Gupshup for appointment reminders, lab result alerts, OTP delivery."),
        ("Lab Analyzers:", "HL7/ASTM bidirectional interface for 50+ auto-analyzer models."),
        ("Biometric Devices:", "ZKTeco, Realtime, Mantra \u2014 push-based attendance via SDK integration."),
        ("Government:", "e-Hospital (NIC), PMJAY, state health portals via standardized APIs."),
    ]
    for bold_p, text in integrations:
        add_bullet(doc, text, bold_p)

    add_section_divider(doc)
    doc.add_page_break()


def build_ai_roadmap(doc):
    add_heading_h1(doc, "AI/ML Innovation Roadmap", "\U0001F916")

    add_callout(doc,
        "12 AI/ML modules across 4 phases \u2014 from predictive analytics and clinical decision support "
        "to ambient documentation and autonomous agentic workflows. Each module is designed with "
        "PII de-identification and DPDP Act compliance built in from Day 1.",
        H_PURPLE_TINT, "\U0001F9E0", H_AI_PURPLE
    )

    add_stat_row(doc, [
        ("12", "", "AI Modules"),
        ("4", "", "Delivery Phases"),
        ("18", "", "Market Gaps Addressed"),
        ("$2.1B", "", "AI Health Market 2030"),
    ])

    # Phase 1
    add_heading_h2(doc, "Phase 1: Foundation AI (Months 1\u20136)", "\u2460")
    phase1 = [
        {
            "title": "Predictive Bed Occupancy Engine",
            "desc": "ML model analyzing admission patterns, seasonal trends, and historical data to predict bed demand 72 hours ahead. Enables proactive staffing and resource allocation.",
            "status": "Phase 1"
        },
        {
            "title": "Intelligent Appointment Optimizer",
            "desc": "Reinforcement learning model for dynamic slot allocation, no-show prediction (85%+ accuracy), and overbooking optimization based on specialty-specific patterns.",
            "status": "Phase 1"
        },
        {
            "title": "Smart Inventory Forecasting",
            "desc": "Time-series forecasting (Prophet + LSTM) for consumable demand, drug expiry prediction, and auto-reorder trigger optimization.",
            "status": "Phase 1"
        },
    ]
    for item in phase1:
        add_feature_card(doc, item["title"], item["desc"], item["status"], "Phase 1")

    # Phase 2
    add_heading_h2(doc, "Phase 2: Clinical AI (Months 7\u201312)", "\u2461")
    phase2 = [
        {
            "title": "Clinical Decision Support System (CDSS)",
            "desc": "Real-time clinical alerts during consultation: drug interactions, allergy warnings, contraindications, guideline-based treatment suggestions. Integrated with WHO/ICMR clinical pathways.",
            "status": "Phase 2"
        },
        {
            "title": "Ambient Clinical Documentation",
            "desc": "Voice-to-SOAP AI powered by speech recognition + Claude LLM. Doctor speaks naturally; system generates structured SOAP notes, ICD-10 codes, and prescription drafts.",
            "status": "Phase 2"
        },
        {
            "title": "Diagnostic Imaging AI",
            "desc": "Computer vision models for X-ray, CT, and pathology slide analysis. Triage-level screening with radiologist-in-the-loop verification.",
            "status": "Phase 2"
        },
    ]
    for item in phase2:
        add_feature_card(doc, item["title"], item["desc"], item["status"], "Phase 2")

    # Phase 3
    add_heading_h2(doc, "Phase 3: Revenue & Operations AI (Months 13\u201318)", "\u2462")
    phase3 = [
        {
            "title": "AI-Powered Revenue Cycle Management",
            "desc": "Automated charge capture, coding optimization (ICD-10/CPT), claim denial prediction, and auto-appeal generation. Target: 15\u201320% revenue leakage reduction.",
            "status": "Phase 3"
        },
        {
            "title": "Patient Risk Stratification",
            "desc": "ML-based risk scoring for readmission, sepsis, falls, and deterioration. Real-time alerts to nursing stations with recommended interventions.",
            "status": "Phase 3"
        },
        {
            "title": "Discharge Summary Generator",
            "desc": "Claude-powered automated discharge summary creation from EMR data. De-identified patient context, structured output, doctor review and sign-off.",
            "status": "Phase 3"
        },
    ]
    for item in phase3:
        add_feature_card(doc, item["title"], item["desc"], item["status"], "Phase 3")

    # Phase 4
    add_heading_h2(doc, "Phase 4: Agentic AI (Months 19\u201324)", "\u2463")
    phase4 = [
        {
            "title": "Agentic Clinical Workflows",
            "desc": "Autonomous AI agents that handle multi-step clinical tasks: order entry verification, lab result interpretation chains, treatment plan adjustment suggestions with human-in-the-loop approval gates.",
            "status": "Phase 4"
        },
        {
            "title": "Conversational Patient AI",
            "desc": "WhatsApp/voice-based AI assistant for patients: symptom triage, appointment booking, medication reminders, post-discharge follow-up with escalation to human staff.",
            "status": "Phase 4"
        },
        {
            "title": "Federated Learning Network",
            "desc": "Privacy-preserving multi-hospital ML training. Models improve across the network without sharing raw patient data. DPDP Act compliant by design.",
            "status": "Phase 4"
        },
    ]
    for item in phase4:
        add_feature_card(doc, item["title"], item["desc"], item["status"], "Phase 4")

    add_section_divider(doc)
    doc.add_page_break()


def build_infrastructure(doc):
    add_heading_h1(doc, "Infrastructure & DevOps", "\u2601")

    add_heading_h2(doc, "Deployment Architecture")
    make_table(doc,
        ["Component", "Specification", "Configuration"],
        [
            ["Container Runtime", "Docker (multi-stage builds)", "One service per container, distroless base"],
            ["Orchestration", "AWS EKS (Kubernetes 1.29)", "5 namespaces: platform, clinical, integrations, monitoring, data"],
            ["Auto-Scaling", "HPA + Cluster Autoscaler", "2\u201320 pods (70% CPU/80% mem), 3\u201350 nodes"],
            ["Primary Region", "AWS Mumbai (ap-south-1)", "Multi-AZ deployment, data residency compliance"],
            ["DR Region", "AWS Singapore (ap-southeast-1)", "Active-passive with 15-min RPO, 4-hour RTO"],
            ["Database", "RDS PostgreSQL 16 Multi-AZ", "Read replicas for reporting, PITR enabled"],
            ["Cache", "ElastiCache Redis 7.x Cluster", "3-node cluster, 64GB, role-based TTLs"],
            ["Object Storage", "S3 + CloudFront CDN", "Medical images, documents, static assets"],
            ["DNS & SSL", "Route 53 + ACM", "Wildcard SSL for *.hmssaas.in"],
            ["WAF", "AWS WAF + Shield Advanced", "OWASP rule set, DDoS protection"],
        ],
        col_widths=[3.0, 4.5, 6.5]
    )

    add_heading_h2(doc, "CI/CD Pipeline (8 Stages)")
    stages = [
        ("1. Lint & Format:", "ESLint, Prettier, Checkstyle \u2014 enforced on every commit."),
        ("2. Unit Tests:", "JUnit 5 + Jest, minimum 80% coverage gate."),
        ("3. Security Scan:", "SAST (SonarQube), dependency audit (Snyk), secret detection."),
        ("4. Build:", "Docker multi-stage build, image optimization, layer caching."),
        ("5. Integration Tests:", "Testcontainers with PostgreSQL, Kafka, Redis \u2014 full service mesh."),
        ("6. ECR Push:", "Versioned image push to Amazon ECR with vulnerability scan."),
        ("7. Kubernetes Apply:", "ArgoCD GitOps sync \u2014 declarative deployment manifests."),
        ("8. Smoke Tests:", "E2E health checks, API contract validation, performance baseline."),
    ]
    for bold_p, text in stages:
        add_numbered(doc, bold_p.split(".")[0], text, bold_p.split(". ")[1].rstrip(":"))

    add_heading_h2(doc, "Disaster Recovery & SLA")
    add_stat_row(doc, [
        ("99.9%", "", "Uptime SLA"),
        ("4", "hrs", "RTO"),
        ("15", "min", "RPO"),
        ("< 500", "ms", "API P95 Latency"),
    ])

    add_section_divider(doc)
    doc.add_page_break()


def build_regulatory(doc):
    add_heading_h1(doc, "Regulatory Compliance", "\u2696")

    add_heading_h2(doc, "DPDP Act 2023 (Digital Personal Data Protection)")
    add_callout(doc,
        "India's first comprehensive data protection law. HMS Enterprise implements all required controls "
        "including consent management, data minimization, right to erasure, breach notification workflows, "
        "and data localization (no cross-border transfer of health data).",
        H_GREEN_TINT, "\u2705", "38A169"
    )
    dpdp_items = [
        ("Consent Module:", "Granular consent collection at registration, consultation, and data sharing touchpoints. Consent withdrawal with cascading data handling."),
        ("Data Minimization:", "Collect only what is clinically necessary. Auto-purge non-essential data after retention period."),
        ("Right to Erasure:", "Patient-initiated data deletion workflow with clinical record retention exceptions (7-year minimum)."),
        ("Breach Notification:", "Automated breach detection, CERT-In notification within 6 hours, patient communication templates."),
        ("Data Localization:", "All health data stored in India (AWS Mumbai). No cross-border transfer. DR in Singapore for metadata only."),
    ]
    for bold_p, text in dpdp_items:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "ABDM Compliance (Mandatory 2026)")
    abdm_items = [
        ("ABHA Health ID:", "Linked at patient registration. Bidirectional health record exchange via Health Information Exchange (HIE)."),
        ("FHIR R4 Resources:", "Patient, Encounter, Observation, MedicationRequest, DiagnosticReport \u2014 all mapped to ABDM specifications."),
        ("Health Facility Registry:", "Auto-registration of hospital as HFR-certified facility."),
        ("Unified Health Interface:", "UHI-compliant APIs for appointment booking and teleconsultation."),
    ]
    for bold_p, text in abdm_items:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "NABH 6th Edition (January 2025)")
    add_body(doc,
        "Built-in NABH quality indicator tracking, evidence management, and gap analysis dashboards. "
        "Supports pre-entry, entry, and full accreditation levels with automated compliance scoring."
    )

    add_heading_h2(doc, "Additional Standards")
    make_table(doc,
        ["Standard", "Scope", "Status"],
        [
            ["ISO 27001:2022", "Information Security Management System", "Day 1"],
            ["ISO 27799", "Health Informatics Security", "Day 1"],
            ["OWASP Top 10", "Web Application Security", "Day 1"],
            ["SOC 2 Type II", "Trust Services Criteria", "Phase 2"],
            ["HIPAA", "US Health Data Protection (export markets)", "Phase 3"],
        ],
        col_widths=[3.5, 7.0, 2.5]
    )

    add_section_divider(doc)
    doc.add_page_break()


def build_gtm(doc):
    add_heading_h1(doc, "Go-to-Market Strategy", "\U0001F680")

    add_heading_h2(doc, "Market Opportunity")
    add_body(doc,
        "India's hospital management software market is projected to reach $8.5 billion by 2030, "
        "growing at 18.7% CAGR. With 70,000+ hospitals and only 15% digital adoption, the addressable "
        "market for cloud-native, AI-enabled platforms is massive. HMS Enterprise targets the "
        "underserved mid-market segment (100\u2013500 beds) where legacy vendors have failed to deliver modern solutions."
    )

    add_stat_row(doc, [
        ("$8.5B", "", "Market Size 2030"),
        ("18.7%", "", "CAGR"),
        ("70K+", "", "Hospitals in India"),
        ("15%", "", "Digital Adoption"),
    ])

    add_heading_h2(doc, "Go-to-Market Phases")
    add_heading_h3(doc, "Phase 1: Foundation (Months 1\u20136)")
    phase1_gtm = [
        ("Direct Sales:", "Target 50 mid-size hospitals in Tier 1 cities (Mumbai, Delhi, Bangalore, Chennai, Hyderabad)."),
        ("Free Pilot Program:", "90-day free pilot for 10 anchor hospitals with dedicated onboarding team."),
        ("Content Marketing:", "Publish industry reports, host webinars, build thought leadership."),
    ]
    for bold_p, text in phase1_gtm:
        add_bullet(doc, text, bold_p)

    add_heading_h3(doc, "Phase 2: Expansion (Months 7\u201312)")
    phase2_gtm = [
        ("Channel Partners:", "Onboard 20+ regional resellers and implementation partners."),
        ("Government Tenders:", "Bid for state health department digitization projects (NHM, PMJAY)."),
        ("Integration Marketplace:", "Launch partner ecosystem for lab analyzers, pharmacy chains, insurance companies."),
    ]
    for bold_p, text in phase2_gtm:
        add_bullet(doc, text, bold_p)

    add_heading_h3(doc, "Phase 3: Scale (Months 13\u201324)")
    phase3_gtm = [
        ("International Expansion:", "UAE, Saudi Arabia, East Africa \u2014 markets with similar regulatory requirements and Indian diaspora hospitals."),
        ("AI Differentiation:", "Launch CDSS, ambient AI, and revenue cycle AI as premium differentiators."),
        ("Platform Play:", "Open APIs for third-party developers, marketplace for clinical apps and integrations."),
    ]
    for bold_p, text in phase3_gtm:
        add_bullet(doc, text, bold_p)

    add_section_divider(doc)
    doc.add_page_break()


def build_revenue_model(doc):
    add_heading_h1(doc, "Revenue & Pricing Model", "\U0001F4B5")

    add_heading_h2(doc, "Tiered SaaS Pricing")
    make_table(doc,
        ["Feature", "Starter", "Professional", "Enterprise"],
        [
            ["Monthly Price", "$999\u2013$2,500", "$5,000\u2013$15,000", "$15,000\u2013$25,000"],
            ["Bed Capacity", "20\u2013100 beds", "100\u2013500 beds", "500+ beds"],
            ["Core Modules (6)", "Included", "Included", "Included"],
            ["Advanced Modules", "Add-on", "Included", "Included"],
            ["AI/ML Features", "Not Available", "Basic AI", "Full AI Suite"],
            ["API Rate Limit", "500 req/min", "2,000 req/min", "5,000 req/min"],
            ["Support", "Email (48h SLA)", "Priority (4h SLA)", "Dedicated CSM (1h SLA)"],
            ["Data Retention", "3 years", "5 years", "7 years"],
            ["Custom Branding", "Not Available", "Partial", "Full"],
            ["SLA", "99.5%", "99.9%", "99.95%"],
            ["Onboarding", "Self-service", "Assisted (2 weeks)", "White-glove (4 weeks)"],
        ],
        col_widths=[3.5, 3.0, 3.5, 3.5]
    )

    add_heading_h2(doc, "Revenue Streams")
    streams = [
        ("Subscription Revenue:", "Core SaaS platform fees \u2014 70% of total revenue. Recurring monthly/annual billing."),
        ("AI Add-on Revenue:", "Premium AI modules (CDSS, ambient AI, RCM AI) \u2014 15% of revenue. Per-module pricing."),
        ("Implementation Services:", "One-time onboarding, data migration, custom configuration \u2014 10% of revenue."),
        ("Marketplace Commissions:", "Third-party app integrations, lab analyzer partnerships \u2014 5% of revenue."),
    ]
    for bold_p, text in streams:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "Financial Projections")
    make_table(doc,
        ["Metric", "Year 1", "Year 2", "Year 3", "Year 5"],
        [
            ["Hospitals Onboarded", "25", "100", "350", "1,200"],
            ["ARR", "$1.5M", "$8M", "$28M", "$120M"],
            ["MRR Growth", "15%", "18%", "12%", "8%"],
            ["Gross Margin", "65%", "72%", "78%", "82%"],
            ["CAC Payback", "14 months", "10 months", "7 months", "5 months"],
            ["Net Revenue Retention", "105%", "115%", "125%", "130%"],
            ["Churn Rate", "5%", "3.5%", "2%", "1.5%"],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 2.5]
    )

    add_section_divider(doc)
    doc.add_page_break()


def build_implementation_roadmap(doc):
    add_heading_h1(doc, "Implementation Roadmap", "\U0001F4C5")

    add_heading_h2(doc, "54-Week Delivery Timeline")
    add_body(doc,
        "The platform is delivered across 4 major phases over 54 weeks (13 months), with each phase "
        "culminating in a production-ready release. The phased approach ensures continuous value delivery "
        "while managing complexity and risk."
    )

    make_table(doc,
        ["Phase", "Duration", "Deliverables", "Key Milestones"],
        [
            ["Phase 1: Core Platform", "Weeks 1\u201314", "OPD, Pharmacy, Lab, Finance, IPD, Scheduling", "MVP launch with 3 pilot hospitals"],
            ["Phase 2: Extended Modules", "Weeks 15\u201328", "HR, Inventory, BI, Patient Portal, Admin", "Feature-complete platform release"],
            ["Phase 3: AI & Intelligence", "Weeks 29\u201342", "CDSS, Ambient AI, Predictive Analytics", "AI-powered clinical features live"],
            ["Phase 4: Scale & Optimize", "Weeks 43\u201354", "RCM AI, Agentic Workflows, Marketplace", "Full platform with 50+ tenants"],
        ],
        col_widths=[3.5, 2.5, 5.0, 4.5]
    )

    add_heading_h2(doc, "Delivery Milestones")
    milestones = [
        ("Week 4:", "Architecture finalized, development environment ready, CI/CD pipeline operational."),
        ("Week 8:", "OPD + Pharmacy modules feature-complete. Internal testing begins."),
        ("Week 14:", "Phase 1 release \u2014 core modules deployed to 3 pilot hospitals."),
        ("Week 20:", "Lab + Finance modules live. ABDM integration testing begins."),
        ("Week 28:", "Full 12-module platform released. Channel partner onboarding starts."),
        ("Week 36:", "CDSS and ambient AI in beta with 5 hospitals."),
        ("Week 42:", "AI modules GA release. Revenue cycle AI in beta."),
        ("Week 48:", "Marketplace launched. Agentic workflow beta begins."),
        ("Week 54:", "Full platform GA. 50+ hospital target achieved."),
    ]
    for bold_p, text in milestones:
        add_bullet(doc, text, bold_p)

    add_heading_h2(doc, "Team Structure")
    make_table(doc,
        ["Role", "Count", "Responsibility"],
        [
            ["Product Owner", "1", "Vision, roadmap, stakeholder management"],
            ["Tech Lead / Architect", "1", "Architecture decisions, code quality, technical guidance"],
            ["Backend Engineers", "4", "Spring Boot microservices, Kafka, PostgreSQL"],
            ["Frontend Engineers", "3", "React, TypeScript, PWA, responsive design"],
            ["AI/ML Engineers", "2", "Python, FastAPI, Claude integration, ML models"],
            ["QA Engineers", "2", "Automated testing, integration testing, performance testing"],
            ["DevOps Engineer", "1", "AWS, Kubernetes, CI/CD, monitoring"],
            ["UI/UX Designer", "1", "Design system, wireframes, user research"],
            ["Project Manager", "1", "Sprint planning, delivery tracking, risk management"],
        ],
        col_widths=[3.5, 1.5, 9.0]
    )

    add_section_divider(doc)
    doc.add_page_break()


def build_appendices(doc):
    add_heading_h1(doc, "Appendices & References", "\U0001F4DA")

    add_heading_h2(doc, "Glossary")
    glossary = [
        ("ABDM", "Ayushman Bharat Digital Mission \u2014 India's national health digitization program."),
        ("ABHA", "Ayushman Bharat Health Account \u2014 14-digit unique health ID."),
        ("CDSS", "Clinical Decision Support System \u2014 real-time clinical alerts and recommendations."),
        ("DPDP Act", "Digital Personal Data Protection Act, 2023 \u2014 India's data protection law."),
        ("FHIR R4", "Fast Healthcare Interoperability Resources \u2014 HL7 standard for health data exchange."),
        ("HIS/HMS", "Hospital Information System / Hospital Management System."),
        ("MAR", "Medication Administration Record \u2014 tracks drug administration to inpatients."),
        ("NABH", "National Accreditation Board for Hospitals and Healthcare Providers."),
        ("RCM", "Revenue Cycle Management \u2014 financial process from patient encounter to payment."),
        ("SOAP", "Subjective, Objective, Assessment, Plan \u2014 standard clinical documentation format."),
    ]
    make_table(doc,
        ["Abbreviation", "Definition"],
        [[abbr, defn] for abbr, defn in glossary],
        col_widths=[3.0, 13.0]
    )

    add_heading_h2(doc, "Document History")
    make_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "16 March 2026", AUTHOR, "Initial release \u2014 complete product document"],
        ],
        col_widths=[2.0, 3.0, 4.0, 6.0]
    )

    add_heading_h2(doc, "References")
    refs = [
        "NASSCOM \u2014 India HealthTech Market Report 2025",
        "NITI Aayog \u2014 Digital Health Ecosystem Framework",
        "Ministry of Health \u2014 ABDM Sandbox Documentation v3.0",
        "MeitY \u2014 Digital Personal Data Protection Act, 2023",
        "NABH \u2014 6th Edition Accreditation Standards (January 2025)",
        "WHO \u2014 Digital Health Technical Guidelines 2024",
        "HL7 International \u2014 FHIR R4 Implementation Guide",
        "Gartner \u2014 Market Guide for Clinical Decision Support 2025",
        "McKinsey \u2014 AI in Healthcare: Opportunities & Challenges 2025",
        f"{COMPANY_FULL} \u2014 HMS Enterprise Technical System Design v1.0",
        f"{COMPANY_FULL} \u2014 HMS Enterprise API Specification v1.0",
        f"{COMPANY_FULL} \u2014 HMS Enterprise SOW Template v1.0",
    ]
    for i, ref in enumerate(refs, 1):
        add_numbered(doc, str(i), ref)


def build_back_page(doc):
    """Branded closing page."""
    doc.add_page_break()

    for _ in range(6):
        sp = doc.add_paragraph()
        add_run(sp, " ", size=14)

    # Navy block
    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_shading(p, H_DARK_NAVY)
        add_run(p, " ", size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p, H_DARK_NAVY)
    p.paragraph_format.space_before = Pt(10)
    add_run(p, COMPANY_FULL.upper(), size=18, bold=True, color=GOLD)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p2, H_DARK_NAVY)
    add_run(p2, "Technology  \u2022  Digital Transformation  \u2022  Healthcare Innovation", size=9, color=RGBColor(160, 174, 192))

    # Gold bar
    g = doc.add_paragraph()
    g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(g, H_GOLD)
    add_run(g, " ", size=3)

    # Contact block
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p3, H_DARK_NAVY)
    p3.paragraph_format.space_before = Pt(8)
    add_run(p3, f"Author: {AUTHOR}", size=10, color=WHITE, light=True)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p4, H_DARK_NAVY)
    add_run(p4, f"Document ID: {DOC_ID}  |  Version {VERSION}  |  {DOC_DATE}", size=8, color=GRAY)

    for _ in range(2):
        p = doc.add_paragraph()
        set_paragraph_shading(p, H_DARK_NAVY)
        add_run(p, " ", size=10)

    # Classification
    doc.add_paragraph()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(cp, H_GOLD_TINT)
    set_paragraph_borders_full(cp, H_GOLD, "4")
    add_run(cp, "  CONFIDENTIAL  ", size=10, bold=True, color=GOLD)

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Pt(12)
    add_run(p5, f"\u00A9 2026 {COMPANY_FULL}. All rights reserved.", size=8, color=GRAY, light=True)


# ═══════════════════════════════════════════════════════════════════
#  MAIN GENERATOR
# ═══════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')

    # Build all sections
    print("Building cover page...")
    build_cover_page(doc)

    print("Adding header & footer...")
    add_page_header(doc)
    add_page_footer(doc)

    print("Building table of contents...")
    build_toc(doc)

    print("Building executive summary...")
    build_executive_summary(doc)

    print("Building product vision & strategy...")
    build_product_vision(doc)

    print("Building platform architecture...")
    build_architecture(doc)

    print("Building module specifications (12 modules)...")
    build_modules(doc)

    print("Building RBAC & security framework...")
    build_rbac_security(doc)

    print("Building API design & integration layer...")
    build_api_design(doc)

    print("Building AI/ML innovation roadmap...")
    build_ai_roadmap(doc)

    print("Building infrastructure & DevOps...")
    build_infrastructure(doc)

    print("Building regulatory compliance...")
    build_regulatory(doc)

    print("Building go-to-market strategy...")
    build_gtm(doc)

    print("Building revenue & pricing model...")
    build_revenue_model(doc)

    print("Building implementation roadmap...")
    build_implementation_roadmap(doc)

    print("Building appendices & references...")
    build_appendices(doc)

    print("Building back page...")
    build_back_page(doc)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved to: {OUTPUT_PATH}")
    print("Done!")


if __name__ == "__main__":
    main()
