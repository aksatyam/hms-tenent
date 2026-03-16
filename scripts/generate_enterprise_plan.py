#!/usr/bin/env python3
"""
HMS Enterprise Detailed Build Plan
Enterprise-grade technical specification document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

OUTPUT_PATH = "/Users/aksatyam/PRODUCT DESIGN/HMS/HMS_Enterprise_Build_Plan.docx"

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge])
            element.set(qn('w:sz'), '4')
            element.set(qn('w:color'), '003366')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.15

# ============================================
# TITLE PAGE
# ============================================

# Company Logo Area
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("TECHNICAL SPECIFICATION")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)
run.font.lt_escapes = 'single'

doc.add_paragraph()

# Main Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("HMS PRO")
run.font.size = Pt(52)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Hospital Management System")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()

# Document Type
doc_type = doc.add_paragraph()
doc_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doc_type.add_run("ENTERPRISE BUILD PLAN")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(51, 51, 51)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Version Info Box
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ('Document ID', 'HMS-TSP-2026-001'),
    ('Version', '1.0'),
    ('Classification', 'Confidential - Internal Use Only'),
    ('Target Release', 'Q4 2026'),
    ('Prepared By', 'TechDigital WishTree'),
    ('Date', datetime.now().strftime('%B %d, %Y'))
]

for i, (label, value) in enumerate(info_data):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
    set_cell_shading(table.rows[i].cells[0], 'E8F4F8')
    table.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph()

# Approval Signatures
para = doc.add_paragraph()
para.add_run("APPROVALS")
para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.runs[0].font.bold = True
para.runs[0].font.size = Pt(14)

# ============================================
# TABLE OF CONTENTS
# ============================================

doc.add_page_break()

toc = doc.add_paragraph()
run = toc.add_run("TABLE OF CONTENTS")
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

toc_items = [
    ("1.", "Document Control", "8"),
    ("2.", "Executive Summary", "9"),
    ("3.", "Product Overview", "10"),
    ("4.", "Technical Architecture", "12"),
    ("5.", "Multi-Tenancy Architecture", "15"),
    ("6.", "RBAC & Security Framework", "18"),
    ("7.", "Functional Modules", "22"),
    ("8.", "API Specifications", "35"),
    ("9.", "Database Design", "40"),
    ("10.", "Integration Requirements", "45"),
    ("11.", "Infrastructure & Deployment", "50"),
    ("12.", "AI & Advanced Features", "55"),
    ("13.", "Delivery Phases", "58"),
    ("14.", "Risk Management", "62"),
    ("15.", "Appendices", "65")
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(f"{num}  {title}")
    p.add_run(f".".rjust(60 - len(f"{num}  {title}"), "."))
    p.add_run(f"{page}")

# ============================================
# 1. DOCUMENT CONTROL
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("1. DOCUMENT CONTROL")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# Revision History
p = doc.add_paragraph()
p.add_run("1.1 Revision History").font.bold = True

table1 = doc.add_table(rows=4, cols=5)
table1.style = 'Table Grid'

headers = ['Version', 'Date', 'Author', 'Changes', 'Status']
for i, h in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

revisions = [
    ('0.1', '2026-02-15', 'Ashish Kumar Satyam', 'Initial draft', 'Draft'),
    ('0.5', '2026-03-01', 'Ashish Kumar Satyam', 'Added technical architecture', 'Review'),
    ('1.0', '2026-03-12', 'Ashish Kumar Satyam', 'Final version for approval', 'Approved')
]

for row_idx, (ver, date, author, changes, status) in enumerate(revisions, start=1):
    table1.rows[row_idx].cells[0].text = ver
    table1.rows[row_idx].cells[1].text = date
    table1.rows[row_idx].cells[2].text = author
    table1.rows[row_idx].cells[3].text = changes
    table1.rows[row_idx].cells[4].text = status

# ============================================
# 2. EXECUTIVE SUMMARY
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("2. EXECUTIVE SUMMARY")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("2.1 Purpose").font.bold = True
p = doc.add_paragraph()
p.add_run("This document provides a comprehensive technical specification and build plan for the HMS Pro (Hospital Management System) - an enterprise-grade, multi-tenant SaaS platform designed specifically for the Tier-2 and Tier-3 Indian healthcare market.")

p = doc.add_paragraph()
p.add_run("2.2 Scope").font.bold = True
p = doc.add_paragraph()
p.add_run("HMS Pro encompasses 12 integrated modules spanning clinical operations, administrative management, and platform services. The platform is architected for scalability to support 50+ concurrent hospital clients with full data isolation and enterprise-grade security compliance.")

p = doc.add_paragraph()
p.add_run("2.3 Key Differentiators").font.bold = True

differentiators = [
    "Multi-Tenant SaaS Architecture: Schema-per-tenant PostgreSQL with Flyway-managed migrations",
    "Event-Driven Real-Time Integration: Apache Kafka event bus with idempotent consumers",
    "Indian Regulatory Compliance: DISHA Act, ABDM/ABHA integration, GST-compliant invoicing",
    "Enterprise RBAC: 6-tier role hierarchy with 200+ permissions and field-level access control",
    "Clinical Excellence: SOAP-based EMR, ICD-10 diagnosis, drug interaction checker (Rxnorm)",
    "AI-Enabled: Claude API integration for discharge summary generation"
]

for diff in differentiators:
    doc.add_paragraph(diff, style='List Bullet')

# ============================================
# 3. PRODUCT OVERVIEW
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("3. PRODUCT OVERVIEW")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("3.1 Problem Statement").font.bold = True
p = doc.add_paragraph()
p.add_run("Healthcare providers in Tier-2 and Tier-3 India face a structural software gap. Enterprise HMS platforms (Attune, Oracle Health) are priced and architected for large corporate chains, making them inaccessible to the majority of India's hospital ecosystem. Affordable solutions (SoftClinic, KnowIT) cover basic admin but lack the clinical depth, multi-branch control, and access governance that growing hospitals need.")

p = doc.add_paragraph()
p.add_run("3.2 Product Vision").font.bold = True
p = doc.add_paragraph()
p.add_run("HMS Pro occupies the white space: a comprehensive, affordable SaaS platform that scales from a solo-doctor clinic to a multi-branch hospital chain. Key strategic pillars include:")

pillars = [
    "Tier-2/3 City Focus — affordable SaaS pricing, vernacular UI support, low-bandwidth optimization",
    "Multi-Tenant Architecture — single platform, unlimited hospital clients, full data isolation",
    "Multi-Branch Management — hospital chain with multiple branches managed from one console",
    "Customizable Multi-Level RBAC — field-level access control for every clinical and admin role",
    "Clinical Management First — full OPD/IPD, EMR, ICU, OT, nursing notes",
    "AI-Enabled Analytics — bed utilization, revenue cycle, clinical outcome trends"
]

for pillar in pillars:
    doc.add_paragraph(pillar, style='List Bullet')

p = doc.add_paragraph()
p.add_run("3.3 Target Market").font.bold = True

table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Table Grid'

market = [
    ('Primary Target', 'Tier-2 and Tier-3 cities in India'),
    ('Secondary Target', 'Multi-branch hospital chains in metro cities'),
    ('Market Size', '50,000+ hospitals and 100,000+ clinics'),
    ('Pricing Model', 'SaaS subscription (per bed/per month)')
]

for i, (label, value) in enumerate(market):
    table2.rows[i].cells[0].text = label
    run = table2.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    set_cell_shading(table2.rows[i].cells[0], 'F0F0F0')
    table2.rows[i].cells[1].text = value

# ============================================
# 4. TECHNICAL ARCHITECTURE
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("4. TECHNICAL ARCHITECTURE")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("4.1 Architectural Style").font.bold = True
p = doc.add_paragraph()
p.add_run("HMS Pro is designed as a cloud-native, multi-tenant SaaS platform built on a modular microservices architecture. Each functional domain (Clinical, Billing, Pharmacy, Lab, RBAC, Tenant Management) is an independently deployable service communicating over well-defined APIs and an asynchronous event bus.")

p = doc.add_paragraph()
p.add_run("4.2 Technology Stack").font.bold = True

tech_table = doc.add_table(rows=8, cols=2)
tech_table.style = 'Table Grid'

tech_stack = [
    ('Frontend', 'React 18 + TypeScript, Ant Design Pro, Vite, PWA'),
    ('API Gateway', 'Spring Cloud Gateway, JWT + Refresh Tokens, Resilience4j'),
    ('Backend', 'Spring Boot 3.x, Java 21, DDD + CQRS Pattern'),
    ('Real-time Messaging', 'Apache Kafka, Socket.IO, Redis Pub/Sub'),
    ('Database', 'PostgreSQL 15, Redis 7, Elasticsearch 8'),
    ('Object Storage', 'MinIO/S3 for documents and reports'),
    ('Infrastructure', 'Kubernetes (EKS), Terraform IaC, GitHub Actions'),
    ('Authentication', 'Keycloak Identity Provider, JWT, MFA')
]

for i, (category, tech) in enumerate(tech_stack):
    tech_table.rows[i].cells[0].text = category
    run = tech_table.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    set_cell_shading(tech_table.rows[i].cells[0], '003366')
    tech_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    tech_table.rows[i].cells[1].text = tech

p = doc.add_paragraph()
p.add_run("4.3 High-Level Architecture Layers").font.bold = True

layers = [
    ("Presentation Layer", "React 18 + TypeScript with Ant Design Pro components"),
    ("API Gateway Layer", "Spring Cloud Gateway for authentication, routing, rate limiting"),
    ("Application Layer", "Spring Boot microservices with DDD + CQRS pattern"),
    ("Event Bus Layer", "Apache Kafka for real-time cross-module communication"),
    ("Data Layer", "PostgreSQL (primary), Redis (caching), Elasticsearch (analytics)")
]

for layer, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"{layer}: ")
    run.font.bold = True
    p.add_run(desc)

p = doc.add_paragraph()
p.add_run("4.4 Key Architectural Decisions").font.bold = True

decisions = [
    "Schema-per-tenant PostgreSQL for complete data isolation",
    "Event-driven architecture using Apache Kafka for loose coupling",
    "CQRS pattern for read/write separation and performance optimization",
    "Redis caching for session management and frequently accessed data",
    "Elasticsearch for full-text search and analytics workloads"
]

for decision in decisions:
    doc.add_paragraph(decision, style='List Bullet')

# ============================================
# 5. MULTI-TENANCY ARCHITECTURE
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("5. MULTI-TENANCY ARCHITECTURE")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("5.1 Tenancy Tiers").font.bold = True
p = doc.add_paragraph()
p.add_run("The platform uses a three-tier tenancy model based on client size and SLA requirements:")

table3 = doc.add_table(rows=4, cols=4)
table3.style = 'Table Grid'

headers = ['Tier', 'Features', 'SLA', 'Pricing']
for i, h in enumerate(headers):
    cell = table3.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

tiers = [
    ('Starter', 'Single branch, 50 beds, 10 users', '99.5%', '₹15,000/month'),
    ('Professional', '5 branches, 200 beds, 50 users', '99.9%', '₹50,000/month'),
    ('Enterprise', 'Unlimited branches, custom SLA', '99.99%', 'Custom')
]

for row_idx, (tier, features, sla, pricing) in enumerate(tiers, start=1):
    table3.rows[row_idx].cells[0].text = tier
    table3.rows[row_idx].cells[1].text = features
    table3.rows[row_idx].cells[2].text = sla
    table3.rows[row_idx].cells[3].text = pricing

p = doc.add_paragraph()
p.add_run("5.2 Tenant Routing Architecture").font.bold = True
p = doc.add_paragraph()
p.add_run("Every request is routed to the correct tenant context before reaching application services:")

routing_steps = [
    "Client accesses platform via tenant subdomain: apollo-indore.hmspro.in",
    "All downstream microservices read tenant_id from context header — never from request body",
    "Database connection pool resolves the correct schema based on tenant_id",
    "Redis cache keys are always prefixed with tenant_id: {tenant_id}:{resource}:{id}"
]

for step in routing_steps:
    doc.add_paragraph(step, style='List Bullet')

p = doc.add_paragraph()
p.add_run("5.3 Tenant Provisioning Flow").font.bold = True

provisioning = [
    "Super admin initiates tenant creation via Admin Console",
    "Tenant Service creates record in master tenant registry",
    "Database provisioner runs: CREATE SCHEMA tenant_{id}",
    "Row-Level Security policies applied",
    "Default admin user created with temporary password",
    "DNS CNAME record auto-created for tenant subdomain",
    "Tenant status set to ACTIVE; provisioning target < 5 minutes"
]

for step in provisioning:
    doc.add_paragraph(step, style='List Bullet')

p = doc.add_paragraph()
p.add_run("5.4 Data Isolation — Defense in Depth").font.bold = True
p = doc.add_paragraph()
p.add_run("Data isolation is enforced at three independent layers:")

isolation = [
    "Application Layer: tenant_id validated on every query",
    "Database Layer: Row-Level Security (RLS) policies",
    "Network Layer: VPC isolation per tenant (Enterprise tier)"
]

for layer in isolation:
    doc.add_paragraph(layer, style='List Bullet')

# ============================================
# 6. RBAC & SECURITY FRAMEWORK
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("6. RBAC & SECURITY FRAMEWORK")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("6.1 Permission Model").font.bold = True
p = doc.add_paragraph()
p.add_run("The RBAC engine uses a three-dimensional permission model: Resource × Action × Field. This enables granular control not possible with traditional role-based systems.")

p = doc.add_paragraph()
p.add_run("6.2 Role Hierarchy (6-Tier)").font.bold = True

table4 = doc.add_table(rows=7, cols=4)
table4.style = 'Table Grid'

headers = ['Level', 'Role', 'Scope', 'Description']
for i, h in enumerate(headers):
    cell = table4.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

roles = [
    ('L0', 'Super Admin', 'Platform', 'Multi-tenant management'),
    ('L1', 'Tenant Admin', 'Organization', 'Full hospital/chain access'),
    ('L2', 'Branch Admin', 'Branch', 'Single branch management'),
    ('L3', 'Department Head', 'Department', 'Department oversight'),
    ('L4', 'Clinical Staff', 'Department', 'Doctor, Nurse, Pharmacist'),
    ('L5', 'Support Staff', 'Assigned', 'Front desk, Billing')
]

for row_idx, (level, role, scope, desc) in enumerate(roles, start=1):
    table4.rows[row_idx].cells[0].text = level
    table4.rows[row_idx].cells[1].text = role
    table4.rows[row_idx].cells[2].text = scope
    table4.rows[row_idx].cells[3].text = desc

p = doc.add_paragraph()
p.add_run("6.3 Field-Level Permission Control").font.bold = True

field_perms = [
    "Nurse role — can VIEW prescription field but cannot EDIT",
    "Intern role — prescription field visible but 'Finalize' button hidden",
    "Billing role — patient demographics visible; clinical notes field hidden",
    "Permission changes take effect within 60 seconds without requiring re-login"
]

for perm in field_perms:
    doc.add_paragraph(perm, style='List Bullet')

p = doc.add_paragraph()
p.add_run("6.4 Authentication & Security").font.bold = True

table5 = doc.add_table(rows=7, cols=2)
table5.style = 'Table Grid'

security = [
    ('Authentication', 'JWT + Refresh tokens, MFA mandatory for L0-L2 roles'),
    ('Session Management', 'Single active session for sensitive roles'),
    ('Failed Login', '5 consecutive failures → account lockout + Admin alert'),
    ('Device Verification', 'OTP re-verification for new device/IP'),
    ('Data Encryption', 'AES-256 for backups, TLS 1.3 in transit'),
    ('Audit Logging', 'Immutable append-only logs, 7-year retention')
]

for i, (feature, impl) in enumerate(security):
    table5.rows[i].cells[0].text = feature
    run = table5.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    set_cell_shading(table5.rows[i].cells[0], 'F0F0F0')
    table5.rows[i].cells[1].text = impl

p = doc.add_paragraph()
p.add_run("6.5 DPDP Act 2023 Compliance").font.bold = True

dpdp = [
    "Explicit patient consent collected at registration",
    "Consent log with timestamp and IP address",
    "Right to erasure: patient data anonymization workflow",
    "Data minimization: API responses return only required fields",
    "Data breach notification: 72-hour workflow to CERT-IN"
]

for item in dpdp:
    doc.add_paragraph(item, style='List Bullet')

# ============================================
# 7. FUNCTIONAL MODULES
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("7. FUNCTIONAL MODULES")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# Module 01: OPD
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.1 Module M01: OPD (Outpatient Department)")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

table6 = doc.add_table(rows=5, cols=3)
table6.style = 'Table Grid'

headers = ['Feature ID', 'Feature', 'Description']
for i, h in enumerate(headers):
    cell = table6.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

opd_features = [
    ('FR-OPD-001', 'Appointment Management', 'Book via front desk, portal, or WhatsApp'),
    ('FR-OPD-002', 'Consultation', 'SOAP format notes, vitals entry, ICD-10'),
    ('FR-OPD-003', 'E-Prescription', 'Drug search, interaction checker, digital signature'),
    ('FR-OPD-004', 'Patient Registration', 'UHID generation, demographic capture')
]

for row_idx, (fid, feature, desc) in enumerate(opd_features, start=1):
    table6.rows[row_idx].cells[0].text = fid
    table6.rows[row_idx].cells[1].text = feature
    table6.rows[row_idx].cells[2].text = desc

# Module 02: Pharmacy
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.2 Module M02: Pharmacy & Drug Management")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

table7 = doc.add_table(rows=5, cols=3)
table7.style = 'Table Grid'

for i, h in enumerate(headers):
    cell = table7.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

pharmacy_features = [
    ('FR-PHM-001', 'FIFO Dispensing', 'First-In-First-Out batch management'),
    ('FR-PHM-002', 'Drug Interaction', 'Rxnorm-based interaction alerts'),
    ('FR-PHM-003', 'Batch Traceability', 'Full traceability from procurement to patient'),
    ('FR-PHM-004', 'Stock Management', 'Automated reorder triggers, inventory tracking')
]

for row_idx, (fid, feature, desc) in enumerate(pharmacy_features, start=1):
    table7.rows[row_idx].cells[0].text = fid
    table7.rows[row_idx].cells[1].text = feature
    table7.rows[row_idx].cells[2].text = desc

# Module 03: Laboratory
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.3 Module M03: Laboratory & Diagnostics")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

table8 = doc.add_table(rows=5, cols=3)
table8.style = 'Table Grid'

for i, h in enumerate(headers):
    cell = table8.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

lab_features = [
    ('FR-LAB-001', 'Sample Management', 'Barcode-tracked from collection to reporting'),
    ('FR-LAB-002', 'Analyzer Integration', 'HL7/ASTM protocol support'),
    ('FR-LAB-003', 'Dual Validation', 'Technical + Medical approval workflow'),
    ('FR-LAB-004', 'Report Delivery', 'PDF with secure S3 URLs, instant delivery')
]

for row_idx, (fid, feature, desc) in enumerate(lab_features, start=1):
    table8.rows[row_idx].cells[0].text = fid
    table8.rows[row_idx].cells[1].text = feature
    table8.rows[row_idx].cells[2].text = desc

# Module 04: IPD
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.4 Module M04: IPD (In-Patient Department)")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

ipd_features = [
    'Bed Management: Live bed availability map, color-coded status',
    'Admission Workflow: One-click admission from OPD referral',
    'Nursing Notes: Shift handover, observation entries, incident reports',
    'Medication Administration Record (MAR): Every dose with nurse ID + timestamp',
    'Doctor Ward Rounds: Structured daily progress notes',
    'Discharge Process: AI-assisted summary, ABHA-linked records'
]

for feature in ipd_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 05: Finance
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.5 Module M05: Finance & Billing")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

billing_features = [
    'Multi-module bill aggregation (OPD + IPD + Pharmacy + Lab)',
    'GST-compliant invoicing (CGST + SGST)',
    'TPA/Insurance pre-authorization workflow',
    'Payment gateway integration (Razorpay/Paytm)',
    'Day-end close automation',
    'Financial reports: P&L, Balance Sheet, Cash Flow'
]

for feature in billing_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 06: HR
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.6 Module M06: HR & Payroll")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

hr_features = [
    'Biometric attendance integration (ZKTeco/Suprema)',
    'Leave Management: CL, SL, EL, Maternity, Paternity',
    'Shift Scheduling: Morning, Evening, Night, Rotational',
    'Payroll Computation: Basic + HRA + TA + Allowances + OT',
    'Deductions: PF (12%+12%), ESIC (0.75%+3.25%), Professional Tax',
    'Employee self-service portal'
]

for feature in hr_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 07: Inventory
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.7 Module M07: Inventory & Procurement")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

inventory_features = [
    'Auto purchase requisition when stock <= reorder level',
    'Multi-level approval workflow (configurable)',
    'Vendor management with performance rating',
    'GRN (Good Received Note) with quality control',
    'Inter-branch transfers',
    'Fixed asset lifecycle management'
]

for feature in inventory_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 08: BI
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.8 Module M08: Business Intelligence & Reports")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

bi_features = [
    'Real-time Kafka event ingestion to Elasticsearch',
    'Nightly ETL aggregation (revenue, OPD volume, bed occupancy)',
    'Executive dashboards: Tenant-wide and branch-level views',
    'REST API for BI tools (Power BI, Tableau, Looker Studio)',
    'Webhook streaming for real-time events',
    'Rate-limited API access'
]

for feature in bi_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 09: Patient Portal
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.9 Module M09: Patient Portal & PWA")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

portal_features = [
    'DISHA-aligned self-service portal',
    'OTP-based authentication',
    'Appointment booking by specialty, branch, doctor',
    'Lab report access with secure PDF download',
    'Online payment processing',
    'Multi-channel notifications (SMS, WhatsApp, Push)'
]

for feature in portal_features:
    doc.add_paragraph(feature, style='List Bullet')

# Module 10: Audit
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("7.10 Module M10: Audit, Compliance & Security")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 102, 153)

audit_features = [
    'Immutable audit log (CREATE/UPDATE/DELETE tracked)',
    '7-year data retention policy',
    'DISHA compliance implementation',
    'AES-256 encrypted backups',
    'OWASP security hardening',
    'Security dashboard with real-time monitoring'
]

for feature in audit_features:
    doc.add_paragraph(feature, style='List Bullet')

# ============================================
# 8. API SPECIFICATIONS
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("8. API SPECIFICATIONS")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("8.1 API Base URL & Versioning").font.bold = True

table9 = doc.add_table(rows=4, cols=2)
table9.style = 'Table Grid'

base_urls = [
    ('Production', 'https://api.hmspro.in/v1'),
    ('Tenant-scoped', 'https://{tenant}.hmspro.in/v1'),
    ('Sandbox', 'https://sandbox.hmspro.in/v1')
]

for i, (env, url) in enumerate(base_urls):
    table9.rows[i].cells[0].text = env
    run = table9.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    table9.rows[i].cells[1].text = url

p = doc.add_paragraph()
p.add_run("8.2 Authentication").font.bold = True
p = doc.add_paragraph()
p.add_run("All endpoints (except /auth/*) require a valid JWT Bearer token. Access token lifespan: 15 minutes | Refresh token: 8 hours.")

p = doc.add_paragraph()
p.add_run("8.3 Standard Response Envelope").font.bold = True

response_format = """
Success (2xx):
{
  "success": true,
  "data": { ... },
  "meta": { "page": 1, "total": 120 }
}

Error (4xx / 5xx):
{
  "success": false,
  "error": { "code": "PATIENT_NOT_FOUND", "message": "...", "field": "patient_id" }
}"""

p = doc.add_paragraph()
run = p.add_run(response_format)
run.font.name = 'Courier New'
run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run("8.4 API Endpoint Summary").font.bold = True

table10 = doc.add_table(rows=12, cols=3)
table10.style = 'Table Grid'

headers = ['Module', 'Endpoints', 'Count']
for i, h in enumerate(headers):
    cell = table10.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

api_summary = [
    ('Authentication', '/auth/login, /auth/refresh, /auth/logout', '8'),
    ('Tenant & RBAC', '/tenants, /branches, /roles, /permissions', '12'),
    ('Patient Management', '/patients, /abha', '10'),
    ('OPD', '/appointments, /consultations, /prescriptions', '15'),
    ('IPD', '/admissions, /beds, /nursing-notes', '12'),
    ('Billing', '/invoices, /payments, /claims', '18'),
    ('Pharmacy', '/drugs, /prescriptions, /dispense', '10'),
    ('Laboratory', '/orders, /samples, /results', '12'),
    ('HR & Staff', '/staff, /attendance, /payroll', '15'),
    ('Reports', '/reports, /exports', '8'),
    ('AI & Notifications', '/ai/*, /notifications/*', '6')
]

for row_idx, (module, endpoints, count) in enumerate(api_summary, start=1):
    table10.rows[row_idx].cells[0].text = module
    table10.rows[row_idx].cells[1].text = endpoints
    table10.rows[row_idx].cells[2].text = count

p = doc.add_paragraph()
p.add_run("Total: 70+ REST API Endpoints").font.bold = True

# ============================================
# 9. DATABASE DESIGN
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("9. DATABASE DESIGN")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("9.1 Database Architecture").font.bold = True

db_table = doc.add_table(rows=5, cols=3)
db_table.style = 'Table Grid'

headers = ['Database', 'Purpose', 'Key Features']
for i, h in enumerate(headers):
    cell = db_table.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

databases = [
    ('PostgreSQL 15', 'Primary Data Store', 'Schema isolation, RLS, JSONB'),
    ('Redis 7', 'Caching & Pub/Sub', 'Sessions, Rate limiting, Pub/Sub'),
    ('Elasticsearch 8', 'Analytics & Search', 'Event indexing, Full-text search'),
    ('MinIO/S3', 'Document Storage', 'Reports, Prescriptions, Images')
]

for row_idx, (db, purpose, features) in enumerate(databases, start=1):
    db_table.rows[row_idx].cells[0].text = db
    db_table.rows[row_idx].cells[1].text = purpose
    db_table.rows[row_idx].cells[2].text = features

p = doc.add_paragraph()
p.add_run("9.2 JSONB Usage Strategy").font.bold = True

jsonb_usage = [
    'consultations.soap_notes — Specialty-specific structured data',
    'consultations.vitals — {bp_systolic, bp_diastolic, pulse, temperature, spo2, weight, height, bmi}',
    'prescriptions.drugs — Array of prescribed medications with interaction flags',
    'invoices.line_items — Flexible billing line items with GST rates'
]

for usage in jsonb_usage:
    doc.add_paragraph(usage, style='List Bullet')

p = doc.add_paragraph()
p.add_run("9.3 Backup & Disaster Recovery").font.bold = True

backup_info = [
    'Daily pg_dump → AES-256 encrypted → S3 (Mumbai + Singapore)',
    'PITR: WAL archiving - restore to any second in last 7 days',
    'RTO: 4 hours, RPO: 15 minutes',
    'Redis snapshot every 15 minutes',
    'Monthly DR drill'
]

for info in backup_info:
    doc.add_paragraph(info, style='List Bullet')

# ============================================
# 10. INTEGRATION REQUIREMENTS
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("10. INTEGRATION REQUIREMENTS")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

table11 = doc.add_table(rows=9, cols=3)
table11.style = 'Table Grid'

headers = ['Integration', 'Technology', 'Purpose']
for i, h in enumerate(headers):
    cell = table11.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

integrations = [
    ('Biometric Devices', 'ZKTeco/Suprema REST API', 'Employee attendance tracking'),
    ('Payment Gateways', 'Razorpay/Paytm', 'Online payments'),
    ('SMS/WhatsApp', 'Twilio/MSG91', 'Patient notifications'),
    ('Email', 'SendGrid/SES', 'Communications'),
    ('BI Tools', 'Power BI, Tableau, Looker Studio', 'Analytics export'),
    ('Lab Analyzers', 'HL7/ASTM protocols', 'Diagnostic equipment'),
    ('ABDM/ABHA', 'NHA Gateway', 'Health ID integration'),
    ('Alerting', 'PagerDuty', 'Incident management')
]

for row_idx, (integ, tech, purpose) in enumerate(integrations, start=1):
    table11.rows[row_idx].cells[0].text = integ
    table11.rows[row_idx].cells[1].text = tech
    table11.rows[row_idx].cells[2].text = purpose

# ============================================
# 11. INFRASTRUCTURE & DEPLOYMENT
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("11. INFRASTRUCTURE & DEPLOYMENT")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("11.1 Kubernetes Cluster Design").font.bold = True

table12 = doc.add_table(rows=5, cols=2)
table12.style = 'Table Grid'

namespaces = [
    ('hms-platform', 'Core services: tenant, auth, rbac, patient'),
    ('hms-clinical', 'Clinical services: clinical, billing, pharmacy, lab'),
    ('hms-integrations', 'External: abdm, notification, ai-service'),
    ('hms-monitoring', 'Observability: Prometheus, Grafana, Loki')
]

for i, (ns, desc) in enumerate(namespaces):
    table12.rows[i].cells[0].text = ns
    run = table12.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    set_cell_shading(table12.rows[i].cells[0], 'F0F0F0')
    table12.rows[i].cells[1].text = desc

p = doc.add_paragraph()
p.add_run("11.2 Auto-Scaling Configuration").font.bold = True

scaling = [
    'HPA: min 2 replicas / max 20 per service; scales at 70% CPU or 80% memory',
    'Cluster Autoscaler: 3 to 50 nodes based on pending pod pressure',
    'Clinical services: Higher minimum replicas (3) for availability',
    'Notification/AI services: Burstable, lower minimum (1) with rapid scale-up'
]

for item in scaling:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run("11.3 CI/CD Pipeline (GitHub Actions + ArgoCD)").font.bold = True

pipeline = [
    'Stage 1: Code Quality - ESLint, Prettier, TypeScript type-check',
    'Stage 2: Unit Tests - Jest with 80% coverage gate',
    'Stage 3: Security Scan - Snyk + Semgrep SAST',
    'Stage 4: Docker Build - Multi-stage Dockerfile, git SHA tagging',
    'Stage 5: Integration Tests - Docker Compose test environment',
    'Stage 6: Push to ECR - Image pushed on merge to main',
    'Stage 7: ArgoCD GitOps - Rolling update with health checks',
    'Stage 8: Smoke Tests - Rollback on failure'
]

for stage in pipeline:
    doc.add_paragraph(stage, style='List Bullet')

# ============================================
# 12. AI & ADVANCED FEATURES
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("12. AI & ADVANCED FEATURES")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("12.1 Claude API Integration").font.bold = True

ai_features = [
    'AI-Assisted Discharge Summary Generation - Reduces documentation time by 60%',
    'Clinical Note Summarization - Auto-generate patient history summaries',
    'ICD-10 Code Suggestions - AI-powered diagnosis recommendations',
    'Drug Interaction Analysis - Advanced interaction checking beyond Rxnorm',
    'Revenue Cycle Analytics - Predict billing bottlenecks and optimization opportunities'
]

for feature in ai_features:
    doc.add_paragraph(feature, style='List Bullet')

p = doc.add_paragraph()
p.add_run("12.2 Kafka Event Architecture").font.bold = True

table13 = doc.add_table(rows=9, cols=3)
table13.style = 'Table Grid'

headers = ['Event Topic', 'Source → Target', 'Purpose']
for i, h in enumerate(headers):
    cell = table13.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

events = [
    ('prescription.created', 'OPD → Pharmacy', 'Trigger Rx processing'),
    ('laborder.created', 'OPD → Laboratory', 'Create sample task'),
    ('drug.dispensed', 'Pharmacy → Finance', 'Post patient charges'),
    ('result.validated', 'Laboratory → Finance', 'Update billing'),
    ('payroll.run', 'HR → Finance', 'Process salary'),
    ('day-close.done', 'Finance', 'Trigger nightly ETL'),
    ('bed.allocated', 'IPD → Finance', 'Initialize room charges'),
    ('appointment.booked', 'Appointments → OPD', 'Patient check-in')
]

for row_idx, (event, flow, purpose) in enumerate(events, start=1):
    table13.rows[row_idx].cells[0].text = event
    table13.rows[row_idx].cells[1].text = flow
    table13.rows[row_idx].cells[2].text = purpose

# ============================================
# 13. DELIVERY PHASES
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("13. DELIVERY PHASES")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("13.1 Phased Delivery Timeline").font.bold = True

table14 = doc.add_table(rows=14, cols=5)
table14.style = 'Table Grid'

headers = ['Phase', 'Name', 'Duration', 'Person-Days', 'Key Deliverables']
for i, h in enumerate(headers):
    cell = table14.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

phases = [
    ('P0', 'Foundation & Design', '4 weeks', '40 PD', 'Architecture, ADRs, SOW'),
    ('P1', 'Multi-Tenant Core', '6 weeks', '60 PD', 'Schema, JWT, MFA, Tenant'),
    ('P2', 'RBAC & Branch', '4 weeks', '40 PD', '6-tier roles, 200+ permissions'),
    ('P3', 'OPD Module', '5 weeks', '52 PD', 'Patient, EMR, E-Rx'),
    ('P4', 'Pharmacy', '5 weeks', '46 PD', 'FIFO, Drug interaction'),
    ('P5', 'Laboratory', '4 weeks', '40 PD', 'Sample management'),
    ('P6', 'Finance', '5 weeks', '50 PD', 'Billing, GST'),
    ('P7', 'HR & Payroll', '5 weeks', '48 PD', 'Attendance, Payroll'),
    ('P8', 'Inventory', '4 weeks', '36 PD', 'Procurement'),
    ('P9', 'BI & Reports', '4 weeks', '40 PD', 'Dashboards'),
    ('P10', 'Audit & Security', '3 weeks', '34 PD', 'Compliance'),
    ('P11', 'E2E QA', '5 weeks', '46 PD', 'Testing'),
    ('P12', 'Go-Live', '3 weeks', '22 PD', 'Deployment')
]

for row_idx, (phase, name, duration, pd, deliverables) in enumerate(phases, start=1):
    table14.rows[row_idx].cells[0].text = phase
    table14.rows[row_idx].cells[1].text = name
    table14.rows[row_idx].cells[2].text = duration
    table14.rows[row_idx].cells[3].text = pd
    table14.rows[row_idx].cells[4].text = deliverables

p = doc.add_paragraph()
p.add_run("Total Duration: 42-46 weeks | Total Person-Days: 550+").font.bold = True

# ============================================
# 14. RISK MANAGEMENT
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("14. RISK MANAGEMENT")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

table15 = doc.add_table(rows=7, cols=4)
table15.style = 'Table Grid'

headers = ['Risk', 'Severity', 'Impact', 'Mitigation']
for i, h in enumerate(headers):
    cell = table15.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

risks = [
    ('DISHA/Healthcare Compliance', 'HIGH', 'Legal, Revenue', 'Checklist at design, DPO from Phase 1'),
    ('Multi-Tenant Performance', 'HIGH', 'User Experience', 'Connection pools, k6 load testing'),
    ('Key Developer Unavailability', 'HIGH', 'Timeline', 'Documentation, cross-training, AI assist'),
    ('Scope Expansion', 'HIGH', 'Budget', 'SOW freeze, Change Control Board'),
    ('Kafka Event Bus Complexity', 'MEDIUM', 'Stability', 'DLQ, idempotent consumers'),
    ('Third-Party API Instability', 'MEDIUM', 'Integrations', 'Circuit breakers, fallback providers')
]

for row_idx, (risk, severity, impact, mitigation) in enumerate(risks, start=1):
    table15.rows[row_idx].cells[0].text = risk
    table15.rows[row_idx].cells[1].text = severity
    table15.rows[row_idx].cells[2].text = impact
    table15.rows[row_idx].cells[3].text = mitigation

# ============================================
# 15. APPENDICES
# ============================================

doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("15. APPENDICES")
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.add_run("Appendix A: Performance Requirements").font.bold = True

perf_table = doc.add_table(rows=7, cols=2)
perf_table.style = 'Table Grid'

perf = [
    ('Platform Uptime', '>99.9% (per calendar month)'),
    ('API Response P95', '<300ms under 500 concurrent users'),
    ('Dashboard Load', '<100ms with Redis cache'),
    ('Real-time Widgets', '<200ms latency'),
    ('Event Propagation', '<1 second'),
    ('Max Tenants', '50 concurrent active tenants')
]

for i, (metric, target) in enumerate(perf):
    perf_table.rows[i].cells[0].text = metric
    run = perf_table.rows[i].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    perf_table.rows[i].cells[1].text = target

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Appendix B: Glossary").font.bold = True
doc.add_paragraph("HMS - Hospital Management System")
doc.add_paragraph("RBAC - Role-Based Access Control")
doc.add_paragraph("CQRS - Command Query Responsibility Segregation")
doc.add_paragraph("DDD - Domain-Driven Design")
doc.add_paragraph("DISHA - Digital Personal Data Protection Act")
doc.add_paragraph("ABDM - Ayushman Bharat Digital Mission")
doc.add_paragraph("ABHA - Ayushman Bharat Health Account")
doc.add_paragraph("FHIR - Fast Healthcare Interoperability Resources")
doc.add_paragraph("UHID - Unique Health Identifier")

# ============================================
# DOCUMENT END
# ============================================

doc.add_page_break()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("END OF DOCUMENT")
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run("TechDigital WishTree | Confidential")
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUTPUT_PATH)
print(f"Successfully generated: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH) / 1024:.2f} KB")