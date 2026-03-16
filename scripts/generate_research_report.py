#!/usr/bin/env python3
"""
HMS Enterprise — Deep Market Research, Gap Analysis & AI/ML Innovation Roadmap
Enterprise-grade DOCX document generator with rich professional theme.
Built by TechDigital WishTree (TGWT)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime

OUTPUT_PATH = "/Users/aksatyam/PRODUCT DESIGN/HMS/docs/research/HMS_Market_Research_Gap_Analysis_AI_Roadmap.docx"

# ─── Company Branding ───
COMPANY_FULL = "TechDigital WishTree"
COMPANY_SHORT = "TGWT"
PRODUCT_NAME = "HMS Enterprise"
DOC_ID = "TGWT-HMS-ENT-2026-RES-001"

# ─── Brand Colors ───
NAVY       = RGBColor(27, 58, 92)    # #1B3A5C
DARK_NAVY  = RGBColor(15, 38, 64)    # #0F2640
TEAL       = RGBColor(13, 115, 119)  # #0D7377
TEAL_LIGHT = RGBColor(16, 185, 129)  # #10B981
GOLD       = RGBColor(196, 154, 42)  # #C49A2A
AI_BLUE    = RGBColor(79, 70, 229)   # #4F46E5
AI_PURPLE  = RGBColor(124, 58, 237)  # #7C3AED
WHITE      = RGBColor(255, 255, 255)
BLACK      = RGBColor(26, 32, 44)    # #1A202C
GRAY       = RGBColor(113, 128, 150) # #718096
LIGHT_GRAY = RGBColor(237, 242, 247) # #EDF2F7
RED        = RGBColor(229, 62, 62)   # #E53E3E
ORANGE     = RGBColor(237, 137, 54)  # #ED8936
GREEN      = RGBColor(56, 161, 105)  # #38A169

# Hex versions for shading
H_NAVY      = "1B3A5C"
H_DARK_NAVY = "0F2640"
H_TEAL      = "0D7377"
H_GOLD      = "C49A2A"
H_AI_BLUE   = "4F46E5"
H_AI_PURPLE = "7C3AED"
H_LIGHT     = "F7FAFC"
H_LIGHTER   = "EDF2F7"
H_WHITE     = "FFFFFF"
H_BLUE_TINT = "EBF8FF"
H_PURPLE_TINT = "FAF5FF"
H_GREEN_TINT = "F0FFF4"
H_GOLD_TINT  = "FFFFF0"
H_RED_TINT   = "FFF5F5"
H_ORANGE_TINT = "FFFAF0"

# ─── Helpers ───
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None, color="D4D4D4", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), sz)
            el.set(qn('w:color'), color)
            borders.append(el)
    tcPr.append(borders)

def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_paragraph_borders(paragraph, bottom_color="1B3A5C", bottom_sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), bottom_sz)
    bottom.set(qn('w:color'), bottom_color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_run(para, text, size=11, bold=False, italic=False, color=BLACK, font_name='Calibri'):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = font_name
    return run

def add_heading_styled(doc, text, level=1):
    """Custom styled heading with rich brand theming."""
    if level == 1:
        # Full-width colored bar header
        bar = doc.add_paragraph()
        bar.paragraph_format.space_before = Pt(24)
        bar.paragraph_format.space_after = Pt(0)
        set_paragraph_shading(bar, H_DARK_NAVY)
        add_run(bar, "  ", size=4)  # top pad

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        set_paragraph_shading(p, H_DARK_NAVY)
        p.paragraph_format.left_indent = Cm(0.5)
        add_run(p, text.upper(), size=20, bold=True, color=WHITE)

        # Gold accent underline bar
        accent = doc.add_paragraph()
        accent.paragraph_format.space_before = Pt(0)
        accent.paragraph_format.space_after = Pt(12)
        set_paragraph_shading(accent, H_GOLD)
        add_run(accent, " ", size=3)
        return p
    elif level == 2:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        # Left border accent + teal shading
        set_paragraph_borders_left(p, H_TEAL, "18")
        set_paragraph_shading(p, H_BLUE_TINT)
        p.paragraph_format.left_indent = Cm(0.4)
        add_run(p, "  " + text, size=14, bold=True, color=NAVY)
        return p
    elif level == 3:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        set_paragraph_borders(p, H_GOLD, "4")
        add_run(p, "\u25C6  ", size=10, color=GOLD)
        add_run(p, text, size=12, bold=True, color=NAVY)
        return p

def set_paragraph_borders_left(paragraph, color_hex, sz="12"):
    """Add a left border accent to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), color_hex)
    left.set(qn('w:space'), '8')
    pBdr.append(left)
    pPr.append(pBdr)

def add_section_divider(doc):
    """Add a decorative divider between major sections."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "\u2500" * 40, size=8, color=LIGHT_GRAY)

def add_stat_row(doc, stats):
    """Add a row of stat boxes (label, value, unit) in a table."""
    t = doc.add_table(rows=2, cols=len(stats))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (val, unit, label) in enumerate(stats):
        # Value cell
        c_top = t.rows[0].cells[i]
        c_top.text = ""
        set_cell_shading(c_top, H_DARK_NAVY)
        p = c_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, val, size=22, bold=True, color=RGBColor(167, 139, 250))  # light purple
        if unit:
            add_run(p, f" {unit}", size=10, color=RGBColor(200, 200, 200))
        # Label cell
        c_bot = t.rows[1].cells[i]
        c_bot.text = ""
        set_cell_shading(c_bot, H_NAVY)
        p2 = c_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(6)
        add_run(p2, label.upper(), size=7, bold=True, color=RGBColor(180, 190, 210))
    doc.add_paragraph()  # spacer
    return t

def add_page_footer(doc):
    """Add branded footer to current section."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    set_paragraph_borders(fp, H_LIGHTER, "2")
    add_run(fp, f"{COMPANY_FULL}", size=7, bold=True, color=NAVY)
    add_run(fp, f"  |  {PRODUCT_NAME}  |  ", size=7, color=GRAY)
    add_run(fp, "CONFIDENTIAL", size=7, bold=True, color=GOLD)

def add_page_header(doc):
    """Add branded header to current section."""
    section = doc.sections[-1]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(hp, f"{DOC_ID}", size=7, color=GRAY)
    add_run(hp, f"  |  {COMPANY_SHORT}", size=7, bold=True, color=TEAL)

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    add_run(p, text, size=10.5, color=RGBColor(45, 55, 72))
    return p

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.2 + indent_level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    marker = "  " * indent_level + "  \u2022  "
    add_run(p, marker, size=10.5, color=TEAL)
    if bold_prefix:
        add_run(p, bold_prefix + " ", size=10.5, bold=True, color=NAVY)
    add_run(p, text, size=10.5, color=RGBColor(45, 55, 72))
    return p

def add_callout(doc, text, bg_hex=H_BLUE_TINT, icon="", border_color=H_TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    set_paragraph_shading(p, bg_hex)
    set_paragraph_borders_left(p, border_color, "16")
    p.paragraph_format.line_spacing = 1.4
    if icon:
        add_run(p, icon + "  ", size=12)
    add_run(p, text, size=10, italic=True, color=RGBColor(45, 55, 72))
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg=H_DARK_NAVY, stripe=True):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=9, bold=True, color=WHITE)
        set_cell_shading(cell, header_bg)
        for edge in ['top', 'bottom', 'left', 'right']:
            set_cell_borders(cell, **{edge: 'single'}, color=header_bg)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            # Color-code check/cross marks
            if val in ("Yes", "Full", "Day 1"):
                add_run(p, val, size=9, bold=True, color=GREEN)
            elif val in ("No", "None", "N/A"):
                add_run(p, val, size=9, bold=True, color=RED)
            elif val in ("Partial", "Basic", "Limited"):
                add_run(p, val, size=9, bold=True, color=ORANGE)
            elif val.startswith("Phase"):
                add_run(p, val, size=9, bold=True, color=AI_PURPLE)
            else:
                add_run(p, val, size=9, color=BLACK)

            if stripe and r_idx % 2 == 1:
                set_cell_shading(cell, H_LIGHT)
            set_cell_borders(cell, bottom='single', color="E2E8F0", sz="2")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


# ═══════════════════════════════════════════════════════
# DOCUMENT START
# ═══════════════════════════════════════════════════════
doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.paragraph_format.line_spacing = 1.3

# ═══════════════════════════════════════════════════════
# COVER PAGE — Rich Enterprise Theme
# ═══════════════════════════════════════════════════════

# Full-width dark navy header bar
for _ in range(2):
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(bar, H_DARK_NAVY)
    add_run(bar, " ", size=4)

# Company name bar
co_bar = doc.add_paragraph()
co_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
co_bar.paragraph_format.space_before = Pt(0)
co_bar.paragraph_format.space_after = Pt(0)
set_paragraph_shading(co_bar, H_DARK_NAVY)
add_run(co_bar, "\n", size=2)
add_run(co_bar, "T E C H D I G I T A L   W I S H T R E E", size=11, bold=True, color=RGBColor(167, 139, 250))
add_run(co_bar, "\n", size=2)

# Gold accent strip
gold_strip = doc.add_paragraph()
gold_strip.paragraph_format.space_before = Pt(0)
gold_strip.paragraph_format.space_after = Pt(0)
set_paragraph_shading(gold_strip, H_GOLD)
add_run(gold_strip, " ", size=3)

doc.add_paragraph()

# Classification badge
cl = doc.add_paragraph()
cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
cl.paragraph_format.left_indent = Cm(4)
cl.paragraph_format.right_indent = Cm(4)
set_paragraph_shading(cl, H_PURPLE_TINT)
add_run(cl, "  CONFIDENTIAL  \u2014  STRATEGIC PLANNING DOCUMENT  ", size=8, bold=True, color=AI_PURPLE)

doc.add_paragraph()
doc.add_paragraph()

# Main Title
t1 = doc.add_paragraph()
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(t1, "HMS", size=56, bold=True, color=NAVY)
add_run(t1, " ENTERPRISE", size=56, bold=True, color=TEAL)

# Decorative line
deco = doc.add_paragraph()
deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
deco.paragraph_format.space_before = Pt(2)
deco.paragraph_format.space_after = Pt(2)
add_run(deco, "\u2500\u2500\u2500\u2500  ", size=12, color=GOLD)
add_run(deco, "\u25C6", size=12, color=GOLD)
add_run(deco, "  \u2500\u2500\u2500\u2500", size=12, color=GOLD)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.paragraph_format.space_before = Pt(4)
add_run(t2, "Hospital Management System", size=16, color=GRAY)

doc.add_paragraph()

# Subtitle in themed box
sub_box = doc.add_paragraph()
sub_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_box.paragraph_format.left_indent = Cm(2.5)
sub_box.paragraph_format.right_indent = Cm(2.5)
set_paragraph_shading(sub_box, H_DARK_NAVY)
sub_box.paragraph_format.space_before = Pt(8)
sub_box.paragraph_format.space_after = Pt(8)
add_run(sub_box, "\n", size=4)
add_run(sub_box, "Deep Market Research\n", size=20, bold=True, color=WHITE)
add_run(sub_box, "Gap Analysis & AI/ML Innovation Roadmap", size=18, bold=True, color=RGBColor(167, 139, 250))
add_run(sub_box, "\n", size=4)

doc.add_paragraph()

# Metadata table — richly themed
meta_table = doc.add_table(rows=7, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Document ID", DOC_ID),
    ("Version", "v1.0 \u2014 Final"),
    ("Classification", "Confidential \u2014 Enterprise SaaS"),
    ("Date", datetime.now().strftime("%B %d, %Y")),
    ("Prepared By", f"{COMPANY_FULL} \u2014 Product Strategy"),
    ("Author", "Ashish Kumar Satyam"),
    ("Status", "\u2705  Approved"),
]
for i, (label, value) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    c0.text = ""
    c1.text = ""
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    add_run(p0, "  " + label, size=9, bold=True, color=WHITE)
    set_cell_shading(c0, H_NAVY)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    add_run(p1, "  " + value, size=9, color=NAVY)
    if i % 2 == 0:
        set_cell_shading(c1, H_LIGHT)
    else:
        set_cell_shading(c1, H_WHITE)
    c0.width = Cm(4.5)
    c1.width = Cm(8)

doc.add_paragraph()

# Tagline with decorative frame
tag_frame = doc.add_paragraph()
tag_frame.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_frame.paragraph_format.left_indent = Cm(3)
tag_frame.paragraph_format.right_indent = Cm(3)
set_paragraph_shading(tag_frame, H_GREEN_TINT)
set_paragraph_borders_left(tag_frame, H_TEAL, "16")
add_run(tag_frame, "  India\u2019s First AI-Native Hospital Management Platform  ", size=11, italic=True, color=TEAL)

doc.add_paragraph()

# Bottom bars
for hex_c in [H_GOLD, H_DARK_NAVY, H_DARK_NAVY]:
    bb = doc.add_paragraph()
    bb.paragraph_format.space_before = Pt(0)
    bb.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(bb, hex_c)
    add_run(bb, " ", size=3 if hex_c == H_GOLD else 4)

# Add header/footer branding
add_page_header(doc)
add_page_footer(doc)

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ═══════════════════════════════════════════════════════
doc.add_page_break()

# TOC header bar
toc_bar = doc.add_paragraph()
toc_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_bar.paragraph_format.space_after = Pt(0)
set_paragraph_shading(toc_bar, H_DARK_NAVY)
add_run(toc_bar, "\n", size=6)
add_run(toc_bar, "TABLE OF CONTENTS", size=18, bold=True, color=WHITE)
add_run(toc_bar, "\n", size=6)

toc_accent = doc.add_paragraph()
toc_accent.paragraph_format.space_before = Pt(0)
toc_accent.paragraph_format.space_after = Pt(0)
set_paragraph_shading(toc_accent, H_GOLD)
add_run(toc_accent, " ", size=2)

doc.add_paragraph()

toc_items = [
    ("01", "Executive Summary", "3"),
    ("02", "Market Landscape & Opportunity", "4"),
    ("03", "Competitive Intelligence Matrix", "7"),
    ("04", "Market Gap Analysis: 18 Unmet Needs", "10"),
    ("05", "AI/ML Capability Roadmap: 12 Modules", "14"),
    ("06", "Emerging Technologies Integration", "19"),
    ("07", "Regulatory Landscape & Compliance", "21"),
    ("08", "Go-to-Market Strategy: AI-First Positioning", "23"),
    ("09", "Revenue Model & Pricing", "25"),
    ("10", "Sources & References", "26"),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), alignment=2, leader=1)
    add_run(p, f"  {num}  ", size=11, bold=True, color=TEAL)
    add_run(p, title, size=11, color=NAVY)
    run = p.add_run(f"\t{page}")
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY


# ═══════════════════════════════════════════════════════
# SECTION 01 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "01  Executive Summary")

add_body(doc, "This document presents a comprehensive market research analysis, competitive intelligence study, gap analysis, and AI/ML innovation roadmap for HMS Enterprise \u2014 a multi-tenant SaaS hospital management platform by TechDigital WishTree")

add_body(doc, "The research identifies 18 critical unmet needs in the Indian HMS market and proposes 12 AI/ML capability modules across four implementation phases to position HMS Enterprise as India\u2019s first AI-native hospital management platform.")

# Key stat boxes
add_stat_row(doc, [
    ("$70B", "", "Global HMS 2033"),
    ("$52.8B", "", "India HealthIT 2030"),
    ("71%", "", "US Hospitals on AI"),
    ("18", "", "Market Gaps Found"),
    ("12", "", "AI/ML Modules"),
])

add_heading_styled(doc, "Key Findings at a Glance", level=2)

# KPI table
kpi_headers = ["Metric", "Value", "Significance"]
kpi_rows = [
    ["Global HMS Market (2033)", "$70 Billion", "7.5% CAGR \u2014 massive growth trajectory"],
    ["India Healthcare IT (2030)", "$52.8 Billion", "16% CAGR \u2014 fastest in Asia-Pacific"],
    ["India Hospital Market (2025)", "$193.4 Billion", "Growing to $364.6B by 2034"],
    ["Undigitized Indian Hospitals", "55-60% (~40,000+)", "Primary target market for HMS Enterprise"],
    ["US Hospitals Using AI", "71%", "Up from 66% in 2023 \u2014 accelerating"],
    ["Ambient AI Market (2033)", "$11.58 Billion", "22% CAGR \u2014 Hindi ambient AI = blue ocean"],
    ["Market Gaps Identified", "18 (4 critical)", "Zero Indian vendor addresses the critical gaps"],
    ["AI/ML Modules Planned", "12 across 4 phases", "First-mover advantage in Indian market"],
]
make_table(doc, kpi_headers, kpi_rows)

add_callout(doc, "The central thesis: No Indian HMS vendor offers AI/ML capabilities. Global leaders (Epic, Oracle Health) only started in 2024-25. HMS Enterprise has a 2-3 year window to establish itself as India\u2019s AI-native HMS leader.", H_PURPLE_TINT, "\u26A1")

# ═══════════════════════════════════════════════════════
# SECTION 02 — MARKET LANDSCAPE
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "02  Market Landscape & Opportunity")

add_body(doc, "The global hospital management system market is valued at $40 billion in 2024 and projected to reach $70 billion by 2033 at 7.5% CAGR. India represents the fastest-growing segment in the Asia-Pacific region, driven by regulatory mandates, insurance digitization, and the Ayushman Bharat Digital Mission.")

add_heading_styled(doc, "2.1  Market Sizing", level=2)

mkt_headers = ["Market Segment", "2024/2025 Value", "Projected Value", "CAGR", "Source"]
mkt_rows = [
    ["Global HMS Software", "$40B (2024)", "$70B by 2033", "7.5%", "Verified Market Reports"],
    ["India Healthcare IT", "\u2014", "$52.8B by 2030", "16%", "Grand View Research"],
    ["India Hospital Market", "$193.4B (2025)", "$364.6B by 2034", "~7%", "IMARC Group"],
    ["India EMR Market", "$0.57B (2025)", "\u2014", "\u2014", "Industry estimates"],
    ["Asia-Pacific HMS", "\u2014", "\u2014", "15%", "Mordor Intelligence"],
    ["Predictive Analytics (HC)", "\u2014", "\u2014", "24.7%", "Fortune BI"],
    ["Ambient Clinical AI", "$1.92B (2024)", "$11.58B by 2033", "22%", "Market Research"],
]
make_table(doc, mkt_headers, mkt_rows)

add_heading_styled(doc, "2.2  Growth Drivers \u2014 India", level=2)

drivers = [
    ("ABDM Mandate (2026):", "AB-PMJAY hospitals must be ABDM-compliant. ABHA ID integration mandatory for government reimbursement."),
    ("NABH 6th Edition (Jan 2025):", "Digital health technology required for accreditation. 100+ hospitals received Digital Health Accreditation."),
    ("DPDP Act 2023:", "India\u2019s primary data protection law (replaced DISHA). Penalties up to \u20B9250 Crore for non-compliance."),
    ("Ayushman Bharat Expansion:", "500M+ beneficiaries. Digital-first claims processing driving HMS adoption."),
    ("Tier-2/3 Hospital Boom:", "68% of new capacity outside metros. These hospitals lack affordable enterprise HMS."),
    ("Insurance Digitization:", "TPA/insurance companies requiring digital claim submissions. Paper hospitals losing empanelment."),
]
for prefix, text in drivers:
    add_bullet(doc, text, bold_prefix=prefix)

add_heading_styled(doc, "2.3  Market Structure \u2014 India HMS Landscape", level=2)

add_body(doc, "India has approximately 70,000+ hospitals across all tiers. The HMS penetration varies dramatically by segment:")

seg_headers = ["Segment", "Market Share", "Count (est.)", "Key Players"]
seg_rows = [
    ["No HMS / Paper + Excel", "55-60%", "~40,000+", "None \u2014 primary target market"],
    ["Basic/Local HMS", "20-25%", "~15,000", "SoftClinic, local vendors"],
    ["Mid-tier HMS", "10-12%", "~7,500", "Practo/Insta, MocDoc, Attune"],
    ["Enterprise HMS", "3-5%", "~2,500", "Oracle Health, custom builds"],
    ["Government (NIC)", "5-8%", "~4,000", "NIC eHospital"],
]
make_table(doc, seg_headers, seg_rows)

add_callout(doc, "HMS Enterprise Sweet Spot: The 55-60% undigitized + 20-25% basic HMS hospitals in Tier-2/3 cities \u2014 approximately 50,000+ hospitals needing affordable enterprise-grade solutions.", H_GOLD_TINT, "\U0001F3AF")

add_heading_styled(doc, "2.4  Global vs. India Technology Gap", level=2)

add_body(doc, "A significant technology gap exists between global HMS leaders and Indian vendors. This gap represents both the challenge and the opportunity for HMS Enterprise:")

gap_headers = ["Capability", "US/Global", "India Mid-Tier", "India Budget", "HMS Enterprise"]
gap_rows = [
    ["Ambient Clinical AI", "Yes (62.6%)", "No", "No", "Phase 3"],
    ["AI Agents (Agentic)", "Yes (Epic/Oracle)", "No", "No", "Phase 2"],
    ["Predictive Analytics", "Full", "Basic", "No", "Phase 1"],
    ["CDSS", "Full (400+ FDA)", "No", "No", "Phase 2"],
    ["Multi-Tenancy", "Limited", "Partial", "No", "Full"],
    ["Event-Driven (Kafka)", "Yes", "No", "No", "Full"],
    ["ABDM/ABHA", "N/A", "Partial", "No", "Day 1"],
    ["Patient Portal PWA", "Yes (MyChart)", "Basic", "No", "Full"],
    ["IoT Integration", "Yes (3,850+/hosp)", "No", "No", "Phase 3"],
    ["HL7/FHIR Interop", "Full", "Basic", "No", "Full"],
    ["6-Tier RBAC", "Yes", "Partial", "No", "Full"],
]
make_table(doc, gap_headers, gap_rows)


# ═══════════════════════════════════════════════════════
# SECTION 03 — COMPETITIVE INTELLIGENCE
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "03  Competitive Intelligence Matrix")

add_body(doc, "Analysis of 8 competing HMS platforms across features, architecture, pricing, and strategic positioning.")

add_heading_styled(doc, "3.1  Feature-by-Feature Comparison", level=2)

comp_headers = ["Feature", "HMS Ent.", "Practo", "MocDoc", "Attune", "KareXpert", "NIC", "SoftClinic", "Oracle"]
comp_rows = [
    ["Multi-Tenancy", "Full", "Partial", "No", "No", "Partial", "No", "No", "Full"],
    ["Event-Driven", "Full", "No", "No", "No", "No", "No", "No", "Yes"],
    ["6-Tier RBAC", "Full", "Partial", "Partial", "Partial", "Partial", "Basic", "No", "Yes"],
    ["OPD Smart Queue", "Full", "Yes", "Yes", "Yes", "Yes", "Partial", "Partial", "Yes"],
    ["Pharmacy FIFO", "Full", "Yes", "Yes", "Yes", "Yes", "Partial", "Partial", "Yes"],
    ["Lab HL7", "Full", "Yes", "Partial", "Yes", "Yes", "Partial", "No", "Yes"],
    ["ABDM/ABHA", "Day 1", "Partial", "Partial", "Partial", "Yes", "Yes", "No", "No"],
    ["GST Finance", "Full", "Yes", "Yes", "Yes", "Yes", "Partial", "Partial", "No"],
    ["BI Dashboards", "Full", "Basic", "Basic", "Basic", "Basic", "No", "No", "Yes"],
    ["Patient Portal", "Full", "Yes", "Partial", "Partial", "Yes", "No", "No", "Yes"],
    ["AI/ML", "12 modules", "No", "No", "No", "Basic", "No", "No", "Yes"],
    ["WhatsApp API", "Full", "Partial", "Partial", "No", "Yes", "No", "No", "No"],
    ["Pricing/yr", "\u20B93-12L", "\u20B95-15L", "\u20B91-50L", "\u20B910-30L", "\u20B98-25L", "Free", "\u20B90.5-2L", "\u20B950L+"],
]
make_table(doc, comp_headers, comp_rows)

add_heading_styled(doc, "3.2  Competitor Weakness Analysis", level=2)

competitors = [
    ("Practo/Insta HMS", [
        "No true multi-tenancy \u2014 shared database model risks data leaks",
        "No event-driven architecture \u2014 batch processing delays",
        "Limited RBAC (3-tier only, no field-level permissions)",
        "Zero AI/ML capabilities",
        "Poor offline support for Tier-2/3 areas",
    ]),
    ("Attune NxG", [
        "On-premise legacy \u2014 expensive infrastructure required",
        "No cloud-native option \u2014 can\u2019t scale dynamically",
        "Priced for corporate chains (\u20B910-30L) \u2014 excludes Tier-2/3",
        "No AI capabilities, no ambient intelligence",
        "Slow implementation (6-12 months typical)",
    ]),
    ("MocDoc", [
        "Feature-rich but complex UI \u2014 high training cost",
        "Monolithic architecture \u2014 difficult to customize",
        "No Kafka/event-driven \u2014 inter-module sync delayed",
        "Pricing opaque and wide-ranging (\u20B91L to \u20B950L+)",
    ]),
    ("NIC eHospital", [
        "Government-only \u2014 unavailable for private hospitals",
        "Slow update cycles \u2014 bureaucratic change management",
        "Poor UX \u2014 designed for compliance, not usability",
        "No analytics, no AI, no patient engagement",
    ]),
]
for comp_name, weaknesses in competitors:
    add_heading_styled(doc, comp_name, level=3)
    for w in weaknesses:
        add_bullet(doc, w)

add_heading_styled(doc, "3.3  HMS Enterprise Competitive Advantages", level=2)

advantages = [
    ("Only Indian HMS with Kafka Event Architecture:", "Real-time cross-module communication (prescription \u2192 pharmacy in <1 sec) vs competitors\u2019 batch/REST approach causing 5-30 second delays."),
    ("Schema-Isolated Multi-Tenancy:", "Each hospital gets its own PostgreSQL schema \u2014 zero data leakage risk. Competitors use shared tables with tenant_id filtering."),
    ("6-Tier RBAC with Field-Level Permissions:", "200+ permissions across 6 hierarchy levels. No Indian competitor offers field-level access control."),
    ("AI-Native from Day 1:", "12 planned AI/ML modules \u2014 no Indian HMS has this. Global leaders only started in 2024-25."),
    ("Enterprise Features at Tier-2/3 Pricing:", "\u20B93-12L/year vs Attune\u2019s \u20B910-30L or Oracle\u2019s \u20B950L+. Milestone payments, EMI, free POC."),
    ("ABDM + NABH + DPDP Triple Compliance:", "Built-in compliance for all three frameworks from Day 1."),
]
for prefix, text in advantages:
    add_bullet(doc, text, bold_prefix=prefix)


# ═══════════════════════════════════════════════════════
# SECTION 04 — GAP ANALYSIS
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "04  Market Gap Analysis: 18 Unmet Needs")

add_body(doc, "Through competitive analysis, user interviews, and global benchmarking, we identified 18 significant unmet needs in the Indian HMS market. These are organized by severity and revenue impact.")

# Critical Gaps
add_heading_styled(doc, "4.1  Critical Gaps (No Indian Vendor Addresses)", level=2)

critical_gaps = [
    ("G01: No AI-Powered Clinical Decision Support",
     "71% of US hospitals use predictive AI, but zero Indian HMS vendors offer CDSS. Indian doctors rely entirely on memory for drug interactions, dosage adjustments, and differential diagnosis.",
     "Drug interaction engine, dosage calculator, allergy cross-check, protocol-based alerts (ICMR), differential diagnosis assistant.",
     "High \u2014 Patient safety + clinical efficiency"),
    ("G02: No Ambient Clinical Intelligence",
     "$1.92B global market growing to $11.58B by 2033 at 22% CAGR. Kaiser Permanente deployed Abridge across 40 hospitals. No Indian HMS offers voice-to-EMR.",
     "Hindi + English ambient listening \u2192 auto-generate SOAP notes, prescriptions, and follow-up instructions from doctor-patient conversations.",
     "High \u2014 Doctor satisfaction + 30% more patients/day"),
    ("G03: No Revenue Cycle AI (Billing Leakage)",
     "AI could cut cost-to-collect by 30-60% (McKinsey). Indian hospitals lose \u20B910-20L/year to billing leakage, unbilled services, and claim denials.",
     "Charge capture AI, auto-coding (ICD-10), claim denial prediction, payment optimization, real-time revenue dashboards.",
     "Very High \u2014 \u20B915-30L/year recoverable per hospital"),
    ("G04: No Predictive Bed & Resource Management",
     "US hospitals use ML to predict admissions 72 hours ahead with 85%+ accuracy. Indian hospitals manage beds via whiteboards and phone calls.",
     "ML-based admission prediction, bed occupancy forecasting, discharge planning optimizer, staff scheduling AI.",
     "High \u2014 5-8% occupancy increase"),
]

gap_table_headers = ["Gap ID & Name", "Problem Statement", "HMS Enterprise Fix", "Revenue Impact"]
gap_table_rows = []
for gid, problem, fix, impact in critical_gaps:
    gap_table_rows.append([gid, problem, fix, impact])
make_table(doc, gap_table_headers, gap_table_rows, header_bg="9B2C2C")

# High-Priority Gaps
add_heading_styled(doc, "4.2  High-Priority Gaps (Partially Addressed by 1-2 Vendors)", level=2)

high_gaps = [
    ("G05: Poor Interoperability", "40% of Indian hospitals face integration challenges (Nature, 2024). No FHIR R4 support.", "FHIR R4 native, HL7 v2 adapters, ABDM HIE, open API marketplace"),
    ("G06: No Pharmacy Demand Forecasting", "\u20B98-15L/year lost to drug expiry. No demand prediction.", "ML demand forecasting using historical dispensing + seasonal patterns + disease trends"),
    ("G07: No Insurance Claim Intelligence", "15-25% claim denial rates. Manual pre-auth takes 2-4 hours.", "AI claim scrubber, denial prediction, auto pre-auth, TPA-specific rules"),
    ("G08: No Patient Risk Stratification", "No Indian HMS identifies high-risk patients proactively.", "Readmission prediction, NEWS2 scoring, sepsis screening, chronic disease risk"),
    ("G09: No Telemedicine Integration", "Most HMS treat telemedicine as bolt-on, not integrated.", "Native video consultation within OPD, e-prescription for teleconsult, auto follow-up"),
    ("G10: No Population Health Analytics", "No cross-facility disease prevalence or outcome analytics.", "Multi-tenant analytics aggregation via Kafka \u2192 Elasticsearch, disease surveillance"),
]
h_headers = ["Gap ID", "Problem", "HMS Fix"]
h_rows = [[g[0], g[1], g[2]] for g in high_gaps]
make_table(doc, h_headers, h_rows, header_bg="C05621")

# Medium Gaps
add_heading_styled(doc, "4.3  Medium-Priority Gaps (Market Differentiators)", level=2)

med_gaps = [
    ("G11: No AI Patient Chatbot", "Patients call hospitals 5-8 times for status queries.", "WhatsApp + Portal AI chatbot \u2014 booking, reports, bills, FAQ, multilingual"),
    ("G12: No Staff Performance Analytics", "Administrators can\u2019t measure doctor/nurse productivity.", "Doctor utilization, nurse workload, department KPIs, automated reports"),
    ("G13: No IoT/Wearable Integration", "3,850+ devices/hospital globally; zero in Indian HMS.", "IoT gateway: bedside monitors \u2192 auto-chart vitals \u2192 anomaly alerts"),
    ("G14: No Intelligent Scheduling", "15-25% no-show rates waste doctor time.", "No-show prediction ML, overbooking optimization, risk-based reminders"),
]
m_rows = [[g[0], g[1], g[2]] for g in med_gaps]
make_table(doc, ["Gap ID", "Problem", "HMS Fix"], m_rows, header_bg="92400E")

# Blue-Sky
add_heading_styled(doc, "4.4  Blue-Sky Opportunities (First-Mover Advantage)", level=2)

blue_gaps = [
    ("G15: NLP Medical Coding", "Inova saved $500K with NLP coding. Indian hospitals code manually.", "Auto-suggest ICD-10/SNOMED from clinical notes, procedure auto-coding"),
    ("G16: Hospital Digital Twin", "Simulate patient flow and capacity before real changes.", "Digital twin using historical Kafka events for what-if scenarios"),
    ("G17: AI Fraud Detection", "Insurance fraud costs \u20B910,000+ Cr/year.", "Anomaly detection, duplicate claim flagging, upcoding alerts"),
    ("G18: Genomics/Precision Medicine", "Genetic testing now <$100/genome; no HMS stores genetic data.", "Genetic data in patient record, pharmacogenomics drug alerts"),
]
b_rows = [[g[0], g[1], g[2]] for g in blue_gaps]
make_table(doc, ["Gap ID", "Problem", "HMS Fix"], b_rows, header_bg="4F46E5")


# ═══════════════════════════════════════════════════════
# SECTION 05 — AI/ML ROADMAP
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "05  AI/ML Capability Roadmap: 12 Modules")

add_body(doc, "Phased rollout of AI/ML features that will make HMS Enterprise India\u2019s first AI-native hospital management platform. Each module is designed to leverage the existing Kafka event-driven architecture for real-time data processing.")

add_heading_styled(doc, "5.0  Implementation Timeline Overview", level=2)

timeline_headers = ["Phase", "Timeline", "Focus Area", "AI Modules", "Prerequisites"]
timeline_rows = [
    ["Phase 1", "Months 6-12", "Foundation AI", "AI-01, AI-02, AI-03", "Kafka events + 6 months of data"],
    ["Phase 2", "Months 12-18", "Clinical AI", "AI-04, AI-05, AI-06, AI-07", "Clinical data accumulation"],
    ["Phase 3", "Months 18-24", "Advanced AI", "AI-08, AI-09, AI-10", "Mature platform + partnerships"],
    ["Phase 4", "Months 24-36", "Frontier AI", "AI-11, AI-12", "Network scale + research"],
]
make_table(doc, timeline_headers, timeline_rows, header_bg=H_AI_PURPLE)

# Phase 1
add_heading_styled(doc, "5.1  Phase 1 \u2014 Foundation AI (Months 6-12)", level=2)

ai_modules_p1 = [
    ("AI-01: Predictive Bed Occupancy", "Fills Gap G04 | IPD + BI Module", [
        "ML model trained on admission/discharge Kafka events",
        "72-hour admission prediction (85%+ accuracy target)",
        "Bed turnover optimization with discharge planning AI",
        "Real-time occupancy forecasting dashboard",
        "Staff scheduling recommendations based on predicted load",
    ], "Python (scikit-learn/XGBoost) \u2192 Kafka consumer \u2192 Redis cache \u2192 React dashboard",
    "Reduce empty bed hours by 20%, increase occupancy 5-8%"),

    ("AI-02: Smart Appointment Scheduler", "Fills Gap G14 | Appointments Module", [
        "No-show prediction model (patient history, demographics, day/time)",
        "Dynamic overbooking based on predicted no-show rate",
        "Wait-time optimization across doctors",
        "Auto-send reminders weighted by no-show risk",
        "Smart waitlist management",
    ], "Gradient Boosted Trees \u2192 Kafka event triggers \u2192 WhatsApp API reminders",
    "Reduce no-shows by 30%, increase doctor utilization 15%"),

    ("AI-03: Pharmacy Demand Forecasting", "Fills Gap G06 | Pharmacy + Inventory", [
        "ML demand prediction per drug per branch per week",
        "Seasonal illness pattern recognition (dengue, flu, etc.)",
        "Auto-generate purchase orders based on predicted demand",
        "Expiry wastage minimization through FIFO + demand matching",
        "Substitution recommendations for stockout scenarios",
    ], "Time-series (Prophet/LSTM) \u2192 Kafka drug.dispensed events \u2192 Inventory auto-PO",
    "Reduce expiry wastage by 40%, reduce stockouts by 60%"),
]

for name, subtitle, features, tech, kpi in ai_modules_p1:
    add_heading_styled(doc, name, level=3)
    p = doc.add_paragraph()
    add_run(p, subtitle, size=9.5, italic=True, color=GRAY)
    for f in features:
        add_bullet(doc, f)
    add_callout(doc, f"Tech Stack: {tech}", H_PURPLE_TINT)
    add_callout(doc, f"Target KPI: {kpi}", H_GREEN_TINT, "\U0001F4CA")

# Phase 2
add_heading_styled(doc, "5.2  Phase 2 \u2014 Clinical AI (Months 12-18)", level=2)

ai_modules_p2 = [
    ("AI-04: Clinical Decision Support (CDSS)", "Fills Gap G01 | OPD + Pharmacy", [
        "Real-time drug interaction alerts (severity graded)",
        "Dosage adjustment based on age, weight, renal/hepatic function",
        "Allergy cross-reference with prescription",
        "Protocol-based treatment suggestions (ICMR guidelines)",
        "Differential diagnosis assistant based on symptoms + vitals",
    ], "Knowledge graph (Neo4j) + rule engine + LLM for differential diagnosis",
    "Prevent 5+ adverse drug events/month, 20% faster diagnosis"),

    ("AI-05: Revenue Cycle Intelligence", "Fills Gap G03 | Finance + Billing", [
        "Charge capture AI \u2014 flag unbilled services from Kafka events",
        "Auto-coding: clinical notes \u2192 ICD-10 suggestions",
        "Claim denial prediction before submission (80% accuracy target)",
        "Payment optimization \u2014 best collection route per patient",
        "Real-time revenue leakage dashboard",
    ], "NLP (clinical notes \u2192 codes) + classification model + Kafka charge events",
    "Recover \u20B915-30L/year per hospital, reduce denials by 40%"),

    ("AI-06: Patient Risk Stratification", "Fills Gap G08 | IPD + OPD", [
        "30-day readmission risk score at discharge",
        "Early Warning Score (NEWS2) auto-calculated from vitals",
        "Sepsis screening from lab results + vitals pattern",
        "Fall risk assessment for IPD patients",
        "Chronic disease progression alerts (diabetes, hypertension)",
    ], "ML ensemble (XGBoost + Logistic Regression) \u2192 real-time Kafka vitals stream \u2192 alert engine",
    "Reduce readmissions by 20%, catch 80% of sepsis cases 4hrs earlier"),

    ("AI-07: Insurance Claim Intelligence", "Fills Gap G07 | Finance", [
        "Pre-auth automation with TPA-specific rule engines",
        "Claim scrubber \u2014 auto-fix errors before submission",
        "Denial pattern analysis per TPA/insurer",
        "Auto-appeal generation for denied claims",
        "Real-time claim status tracking dashboard",
    ], "Rule engine + NLP (claim text analysis) + classification model",
    "Reduce pre-auth time from 4hrs to 15min, reduce denials by 35%"),
]

for name, subtitle, features, tech, kpi in ai_modules_p2:
    add_heading_styled(doc, name, level=3)
    p = doc.add_paragraph()
    add_run(p, subtitle, size=9.5, italic=True, color=GRAY)
    for f in features:
        add_bullet(doc, f)
    add_callout(doc, f"Tech Stack: {tech}", H_PURPLE_TINT)
    add_callout(doc, f"Target KPI: {kpi}", H_GREEN_TINT, "\U0001F4CA")

# Phase 3
add_heading_styled(doc, "5.3  Phase 3 \u2014 Advanced AI (Months 18-24)", level=2)

ai_modules_p3 = [
    ("AI-08: Ambient Clinical Intelligence", "Fills Gap G02 | OPD Module", [
        "Hindi + English voice-to-text from doctor-patient conversations",
        "Auto-generate SOAP notes from conversation",
        "Extract prescriptions, lab orders, follow-up instructions",
        "Doctor review + approve workflow",
        "Integrates with existing OPD EMR module",
    ], "Whisper/IndicWhisper (ASR) + Medical LLM (fine-tuned) + Structured output \u2192 EMR",
    "Reduce documentation time by 70%, see 30% more patients"),

    ("AI-09: AI Agents (Agentic Workflows)", "New Capability | Platform-Wide", [
        "Phone scheduling agent \u2014 handles appointment calls in Hindi/English",
        "Insurance pre-auth agent \u2014 auto-submits and follows up",
        "Patient engagement agent \u2014 reminders, reports, queries via WhatsApp",
        "Admin assistant \u2014 generates reports, answers queries about hospital data",
        "Coding agent \u2014 auto-codes procedures and diagnoses",
    ], "LLM orchestration (Claude/GPT) + tool-use APIs + HMS module connectors",
    "Automate 40% of admin tasks, reduce phone staff by 50%"),

    ("AI-10: NLP Medical Coding Engine", "Fills Gap G15 | Finance", [
        "Clinical notes \u2192 ICD-10 code suggestions (top-3 with confidence)",
        "Procedure auto-coding from operative notes",
        "SNOMED CT mapping for ABDM compliance",
        "DRG optimization for package-based billing",
        "Coding audit \u2014 flag under-coded encounters",
    ], "Fine-tuned medical NLP model + ICD-10/SNOMED knowledge graph",
    "85%+ coding accuracy, recover \u20B95-10L/year in undercoded revenue"),
]

for name, subtitle, features, tech, kpi in ai_modules_p3:
    add_heading_styled(doc, name, level=3)
    p = doc.add_paragraph()
    add_run(p, subtitle, size=9.5, italic=True, color=GRAY)
    for f in features:
        add_bullet(doc, f)
    add_callout(doc, f"Tech Stack: {tech}", H_PURPLE_TINT)
    add_callout(doc, f"Target KPI: {kpi}", H_GREEN_TINT, "\U0001F4CA")

# Phase 4
add_heading_styled(doc, "5.4  Phase 4 \u2014 Frontier AI (Months 24-36)", level=2)

add_heading_styled(doc, "AI-11: Hospital Digital Twin", level=3)
for f in [
    "Virtual simulation of patient flow through departments",
    "\u201CWhat-if\u201D capacity planning for expansion",
    "Bottleneck identification via process mining on Kafka events",
    "Outbreak simulation (bed surge, staff reallocation)",
]:
    add_bullet(doc, f)

add_heading_styled(doc, "AI-12: Population Health & Federated Learning", level=3)
for f in [
    "Cross-tenant disease surveillance (privacy-preserving)",
    "Treatment outcome benchmarking across hospital network",
    "Federated ML \u2014 train models across tenants without sharing data",
    "Genomics-ready patient profiles for pharmacogenomics",
]:
    add_bullet(doc, f)


# ═══════════════════════════════════════════════════════
# SECTION 06 — EMERGING TECHNOLOGIES
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "06  Emerging Technologies Integration")

techs = [
    ("6.1  IoT & Wearable Integration", "Phase 3", [
        ("Market:", "7.4M smart medical devices by 2026 (avg 3,850+/hospital globally)"),
        ("Bedside Monitors:", "Auto-chart vitals (HR, BP, SpO2, temp) into EMR every 5 min"),
        ("Infusion Pumps:", "Medication administration tracking, dose alerts"),
        ("Wearables:", "Patient activity tracking, fall detection alerts"),
        ("Architecture:", "MQTT broker \u2192 Kafka IoT topic \u2192 Vitals auto-charting \u2192 EMR + anomaly alert engine"),
    ]),
    ("6.2  Blockchain for Health Records", "Phase 4", [
        ("Consent Ledger:", "Smart contracts for data sharing consent (DPDP Act compliance)"),
        ("Record Integrity:", "Hash of every EMR update stored on-chain for tamper-proof audit"),
        ("Insurance Claims:", "Transparent claim lifecycle tracking (hospital \u2194 TPA \u2194 insurer)"),
        ("Drug Supply Chain:", "Verify drug authenticity from manufacturer to pharmacy"),
        ("Approach:", "Hyperledger Fabric (permissioned) \u2014 lightweight, not full blockchain"),
    ]),
    ("6.3  Native Telemedicine", "Phase 2", [
        ("Compliance:", "India Telemedicine Practice Guidelines (TPG) 2020"),
        ("Video:", "WebRTC-based consultation embedded within OPD module"),
        ("E-Prescription:", "Digital signature valid for teleconsult prescriptions"),
        ("Billing:", "Teleconsult billing with different tariff structure"),
        ("Recording:", "With patient consent for medico-legal compliance"),
    ]),
    ("6.4  Voice-First EMR / Ambient Clinical Intelligence", "Phase 3", [
        ("Market:", "$1.92B (2024) \u2192 $11.58B (2033) at 22% CAGR"),
        ("Abridge:", "$300M Series E, deployed at Kaiser Permanente (40 hospitals), $600M revenue"),
        ("Microsoft:", "Dragon Copilot / Nuance DAX embedded into Epic/Oracle"),
        ("HMS Opportunity:", "None of these support Hindi. HMS Enterprise can be the first Hindi-capable ambient clinical AI \u2014 massive moat."),
    ]),
]

for title, phase, items in techs:
    add_heading_styled(doc, title, level=2)
    p = doc.add_paragraph()
    add_run(p, f"Timeline: {phase}", size=9.5, italic=True, color=AI_PURPLE)
    for prefix, text in items:
        add_bullet(doc, text, bold_prefix=prefix)


# ═══════════════════════════════════════════════════════
# SECTION 07 — REGULATORY
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "07  Regulatory Landscape & Compliance Roadmap")

regs = [
    ("7.1  ABDM (Ayushman Bharat Digital Mission) \u2014 Mandatory 2026", [
        ("ABHA ID:", "Create, verify, link to patient records. Mandatory for AB-PMJAY hospitals."),
        ("HIP/HIU:", "Health Information Provider (push) and User (pull) flows."),
        ("HFR:", "Health Facility Registry \u2014 every hospital must register."),
        ("PHR App:", "Push records to patient\u2019s digital health locker."),
        ("FHIR R4:", "Required data format for ABDM health record exchange."),
    ], "ABDM/ABHA integration planned as Day-1 feature. Full HIP/HIU/HFR compliance."),

    ("7.2  NABH 6th Edition (January 2025) \u2014 Active Now", [
        ("Digital Health Technology:", "New emphasis; 100+ hospitals received Digital Health Accreditation."),
        ("Digital EMR:", "Paper medical records insufficient for accreditation."),
        ("Medication Safety:", "Drug interaction checks, allergy alerts mandatory."),
        ("Quality Indicators:", "Digital dashboards for NABH quality metrics."),
        ("Audit Trail:", "Every clinical action must be logged and traceable."),
    ], "Hospitals without digital HMS will lose NABH accreditation \u2192 lose insurance empanelment \u2192 lose CGHS/ECHS enrollment."),

    ("7.3  DPDP Act 2023 (Data Protection) \u2014 Enacted", [
        ("Note:", "DISHA was never enacted; DPDP Act 2023 (ratified Aug 2023) is India\u2019s primary data protection law."),
        ("Patient Consent:", "Explicit opt-in for data processing, right to withdraw."),
        ("Data Localization:", "Health data must be stored in India."),
        ("Breach Notification:", "Mandatory notification to Data Protection Board."),
        ("Penalties:", "Up to \u20B9250 Crore for non-compliance."),
    ], "Schema-isolated multi-tenancy + AES-256 encryption + audit trail + consent module = DPDP-ready from Day 1."),

    ("7.4  Upcoming Regulations (2026-2028)", [
        ("AI in Healthcare:", "Expected CDSCO/MoHFW guidelines for AI-based clinical tools."),
        ("Telemedicine Act:", "Comprehensive regulation to replace TPG 2020."),
        ("Health Data Exchange:", "ABDM pushing mandatory FHIR R4 for all hospitals."),
        ("Cybersecurity:", "Healthcare as Critical Information Infrastructure (CERT-In)."),
        ("Drug Traceability:", "Track-and-trace for pharmaceutical supply chain."),
    ], None),
]

for title, items, callout_text in regs:
    add_heading_styled(doc, title, level=2)
    for prefix, text in items:
        add_bullet(doc, text, bold_prefix=prefix)
    if callout_text:
        add_callout(doc, callout_text, H_GOLD_TINT, "\u26A0\uFE0F")


# ═══════════════════════════════════════════════════════
# SECTION 08 — GO-TO-MARKET
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "08  Go-to-Market Strategy: AI-First Positioning")

add_heading_styled(doc, "8.1  Positioning Statement", level=2)

pos = doc.add_paragraph()
pos.paragraph_format.left_indent = Cm(1)
pos.paragraph_format.right_indent = Cm(1)
set_paragraph_shading(pos, H_PURPLE_TINT)
add_run(pos, "For ", size=11, italic=True, color=RGBColor(45, 55, 72))
add_run(pos, "Tier-2/3 hospital groups", size=11, bold=True, color=NAVY)
add_run(pos, " who need enterprise-grade digital operations, ", size=11, italic=True, color=RGBColor(45, 55, 72))
add_run(pos, "HMS Enterprise", size=11, bold=True, color=AI_PURPLE)
add_run(pos, " is the only ", size=11, italic=True, color=RGBColor(45, 55, 72))
add_run(pos, "AI-native, multi-tenant hospital management platform", size=11, bold=True, color=NAVY)
add_run(pos, " that delivers ", size=11, italic=True, color=RGBColor(45, 55, 72))
add_run(pos, "12 AI/ML capabilities", size=11, bold=True, color=AI_PURPLE)
add_run(pos, " at Tier-2/3 pricing with Day-1 ABDM + NABH + DPDP compliance.", size=11, italic=True, color=RGBColor(45, 55, 72))

add_heading_styled(doc, "8.2  Three Key Differentiators", level=2)

diffs = [
    ("\u201CAI That Indian Doctors Actually Need\u201D", "Hindi ambient clinical intelligence, ICMR-aligned CDSS, insurance claim AI for Indian TPAs \u2014 not imported US features. Built for Indian clinical workflows from Day 1."),
    ("\u201CEnterprise Architecture, Tier-2 Pricing\u201D", "Kafka event-driven, schema-isolated multi-tenancy, 6-tier RBAC \u2014 architecture Oracle charges \u20B950L+ for \u2014 starting at \u20B93L/year. Milestone payments, free POC."),
    ("\u201CCompliance Without Compromise\u201D", "ABDM + NABH 6th Edition + DPDP Act \u2014 triple compliance built-in. Get NABH-ready in 16 weeks. Don\u2019t lose accreditation or empanelment."),
]
for i, (title, desc) in enumerate(diffs, 1):
    add_heading_styled(doc, f"Differentiator {i}: {title}", level=3)
    add_body(doc, desc, indent=True)

add_heading_styled(doc, "8.3  Target Segment Prioritization", level=2)

seg_headers = ["Priority", "Segment", "Size (est.)", "Pain Point", "HMS Hook", "Deal Size"]
seg_rows = [
    ["P0", "Tier-2/3 multi-branch (50-300 beds)", "~5,000", "No affordable multi-tenant + NABH", "Schema tenancy + NABH + AI", "\u20B95-12L/yr"],
    ["P1", "Single-location (100-200 beds)", "~15,000", "Paper/Excel, losing NABH", "NABH + GST + patient portal", "\u20B93-6L/yr"],
    ["P1", "Diagnostic chains", "~3,000", "Lab TAT, no barcode tracking", "Lab module + ABDM + AI TAT", "\u20B94-8L/yr"],
    ["P2", "Nursing homes (20-50 beds)", "~30,000", "Basic digitization", "Starter: OPD + Pharma + Finance", "\u20B91.5-3L/yr"],
    ["P3", "Government (NIC migration)", "~2,000", "Poor UX, no analytics", "Modern UX + BI + ABDM + WA", "\u20B98-15L/yr"],
]
make_table(doc, seg_headers, seg_rows)


# ═══════════════════════════════════════════════════════
# SECTION 09 — REVENUE MODEL
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "09  Revenue Model & Pricing Strategy")

add_body(doc, "Three-tier pricing strategy with AI/ML as the premium differentiator. Each tier is designed to match a specific hospital segment\u2019s budget and needs.")

pricing_headers = ["", "Starter", "Professional", "Enterprise AI"]
pricing_rows = [
    ["Annual Price", "\u20B93L/year", "\u20B96-8L/year", "\u20B910-15L/year"],
    ["Target Segment", "Nursing homes, clinics", "Mid-size hospitals", "Hospital chains"],
    ["Modules", "OPD + Pharmacy + Finance", "All 12 modules", "All 12 + AI suite"],
    ["RBAC", "3-tier basic", "6-tier full", "6-tier + field-level"],
    ["BI Dashboards", "Standard reports", "Real-time dashboards", "AI-powered analytics"],
    ["Patient Portal", "No", "Full PWA", "Full PWA + AI chatbot"],
    ["AI/ML Modules", "None", "Phase 1 (3 modules)", "All 12 AI modules"],
    ["Branches", "Single", "Up to 5", "Unlimited"],
    ["ABDM/ABHA", "Yes", "Yes", "Yes"],
    ["Support", "Email", "Email + Phone", "Dedicated + SLA"],
    ["Ambient AI", "No", "No", "Yes (Hindi + English)"],
    ["Implementation", "2-4 weeks", "6-8 weeks", "8-12 weeks"],
]
make_table(doc, pricing_headers, pricing_rows)

add_callout(doc, "Revenue Projections: 50 hospitals on Professional plan = \u20B93-4 Cr/year. 20 hospitals on Enterprise AI = \u20B92-3 Cr/year. Combined Year 1 target: \u20B95-7 Cr ARR.", H_GREEN_TINT, "\U0001F4B0")


# ═══════════════════════════════════════════════════════
# SECTION 10 — SOURCES
# ═══════════════════════════════════════════════════════
doc.add_page_break()

add_heading_styled(doc, "10  Sources & References")

source_categories = [
    ("Market Reports", [
        "Verified Market Reports \u2014 Hospital Management System Software Market 2033",
        "Grand View Research \u2014 India Healthcare Information System Market 2030",
        "IMARC Group \u2014 India Hospital Market Size & Share 2025-2034",
        "Mordor Intelligence \u2014 Asia-Pacific HMS Market Analysis",
        "Fortune Business Insights \u2014 Global HMS Market Forecast",
        "Statista \u2014 Healthcare IT Spending India 2024-2030",
    ]),
    ("AI/ML in Healthcare", [
        "Health Affairs \u2014 AI/Predictive Models in US Hospitals (2024)",
        "McKinsey \u2014 Agentic AI & Touchless Revenue Cycle (2025)",
        "Menlo Ventures \u2014 State of AI in Healthcare 2025",
        "AJMC \u2014 Ambient AI Adoption in US Hospitals (62.6%)",
        "Citeline \u2014 HIMSS 2026: AI Integration, ROI, Clinical Pain Points",
        "FDA \u2014 AI/ML Authorized Medical Devices Database",
        "Nature Medicine \u2014 Clinical Decision Support Systems Review",
    ]),
    ("Ambient AI & Voice", [
        "Abridge \u2014 $300M Series E, Kaiser Permanente deployment",
        "Microsoft \u2014 Dragon Copilot / Nuance DAX Integration",
        "Oracle Health \u2014 Next-Gen EHR with Voice AI (Aug 2025)",
        "Epic \u2014 Native AI Agents Preview at UGM 2025",
        "Suki AI \u2014 Voice Assistant Clinical Documentation",
    ]),
    ("India Regulatory", [
        "NABH \u2014 6th Edition Accreditation Standards (Jan 2025)",
        "NHA \u2014 ABDM Compliance Requirements for AB-PMJAY 2026",
        "DPDP Act 2023 \u2014 Digital Personal Data Protection (Aug 2023)",
        "CERT-In \u2014 Healthcare Critical Infrastructure Guidelines",
        "MoHFW \u2014 Telemedicine Practice Guidelines 2020",
        "ICMR \u2014 Ethical Guidelines for AI in Biomedical Research",
    ]),
    ("Digital Transformation", [
        "Nature (2024) \u2014 Barriers to Digital Transformation in Indian Health Sector",
        "WHO \u2014 Digital Health Strategy 2020-2025",
        "Deloitte \u2014 Smart Hospital: Connected, Efficient, Effective",
        "Accenture \u2014 Healthcare Technology Vision 2025",
        "NASSCOM \u2014 India HealthTech Landscape Report 2024",
    ]),
    ("Competitor Intelligence", [
        "Practo/Insta \u2014 18K+ daily users, 1250+ centers, 22 countries",
        "MocDoc \u2014 Feature matrix, pricing ranges",
        "Attune Technologies \u2014 NxG HMS specifications",
        "KareXpert \u2014 Cloud HMS product details",
        "NIC eHospital \u2014 Government HMS, 1000+ facilities",
        "Oracle Health \u2014 Next-gen EHR announcements (2025)",
        "Epic Systems \u2014 41.3% US market share, AI agents",
        "MEDITECH \u2014 Expanse platform, 11.9% market share",
    ]),
]

for cat_title, sources in source_categories:
    add_heading_styled(doc, cat_title, level=3)
    for idx, src in enumerate(sources, 1):
        add_bullet(doc, src)


# ═══════════════════════════════════════════════════════
# BACK COVER / END PAGE
# ═══════════════════════════════════════════════════════
doc.add_page_break()

doc.add_paragraph()
doc.add_paragraph()

# End page branding bars
for _ in range(2):
    ebar = doc.add_paragraph()
    ebar.paragraph_format.space_before = Pt(0)
    ebar.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(ebar, H_DARK_NAVY)
    add_run(ebar, " ", size=4)

end_bar = doc.add_paragraph()
end_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph_shading(end_bar, H_DARK_NAVY)
end_bar.paragraph_format.space_before = Pt(0)
end_bar.paragraph_format.space_after = Pt(0)
add_run(end_bar, "\n\n", size=6)
add_run(end_bar, "END OF DOCUMENT", size=16, bold=True, color=WHITE)
add_run(end_bar, "\n\n", size=6)

gold_end = doc.add_paragraph()
gold_end.paragraph_format.space_before = Pt(0)
gold_end.paragraph_format.space_after = Pt(0)
set_paragraph_shading(gold_end, H_GOLD)
add_run(gold_end, " ", size=3)

for _ in range(2):
    ebar2 = doc.add_paragraph()
    ebar2.paragraph_format.space_before = Pt(0)
    ebar2.paragraph_format.space_after = Pt(0)
    set_paragraph_shading(ebar2, H_DARK_NAVY)
    add_run(ebar2, " ", size=4)

doc.add_paragraph()

end1 = doc.add_paragraph()
end1.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(end1, PRODUCT_NAME, size=14, bold=True, color=NAVY)

end_deco = doc.add_paragraph()
end_deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(end_deco, "\u2500\u2500\u2500  \u25C6  \u2500\u2500\u2500", size=10, color=GOLD)

end2 = doc.add_paragraph()
end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(end2, "Deep Market Research, Gap Analysis & AI/ML Innovation Roadmap", size=10, color=GRAY)

doc.add_paragraph()

end3 = doc.add_paragraph()
end3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(end3, f"Version 1.0  |  {datetime.now().strftime('%B %Y')}  |  ", size=9, color=GRAY)
add_run(end3, COMPANY_FULL, size=9, bold=True, color=TEAL)

doc.add_paragraph()

end4 = doc.add_paragraph()
end4.alignment = WD_ALIGN_PARAGRAPH.CENTER
end4.paragraph_format.left_indent = Cm(4)
end4.paragraph_format.right_indent = Cm(4)
set_paragraph_shading(end4, H_GOLD_TINT)
add_run(end4, "  CONFIDENTIAL \u2014 For internal strategic planning use only  ", size=9, bold=True, color=GOLD)


# ═══════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"\n\u2705 Document saved to: {OUTPUT_PATH}")
print(f"   Sections: 10")
print(f"   Pages: ~28")
print(f"   Tables: 12")
print(f"   AI/ML Modules: 12")
print(f"   Market Gaps: 18")
