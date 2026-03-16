#!/usr/bin/env python3
"""HMS Detailed Build Plan Generator"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

OUTPUT_PATH = "/Users/aksatyam/PRODUCT DESIGN/HMS/HMS_Detailed_Build_Plan.docx"

def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title Page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("HMS")
run.font.size = Pt(48)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Hospital Management System")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
doc.add_paragraph()

main = doc.add_paragraph()
main.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main.add_run("DETAILED BUILD PLAN")
run.font.size = Pt(32)
run.font.bold = True

doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
run.font.size = Pt(14)
run.font.italic = True

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run("Version 1.0 - Implementation Ready")
run.font.size = Pt(12)

# 1. Executive Summary
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("1. Executive Summary")
run.font.size = Pt(16)
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("Project Overview: ")
run.font.bold = True
p.add_run("HMS is a comprehensive, multi-tenant SaaS platform for hospital groups and clinic chains with 12 integrated modules.")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Scope: ")
run.font.bold = True
p.add_run("Full-stack healthcare platform with clinical, administrative, and platform modules. Built for enterprise scale.")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Timeline: ")
run.font.bold = True
p.add_run("Approximately 42-46 weeks (13 phases)")

# 2. Module Overview
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("2. Module Overview (12 Modules)")
run.font.size = Pt(16)
run.font.bold = True

# Clinical Modules Table
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'

headers = ['Module', 'ID', 'Description']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

clinical_modules = [
    ('OPD', 'M01', 'Outpatient - UHID, EMR, E-Prescription'),
    ('Pharmacy', 'M02', 'Drug Management - FIFO, Batch Trace'),
    ('Laboratory', 'M03', 'Diagnostics - HL7/ASTM Integration'),
    ('IPD', 'M10', 'In-Patient - Bed Mgmt, Nursing')
]

for row_idx, (name, id_, desc) in enumerate(clinical_modules, start=1):
    table.rows[row_idx].cells[0].text = f"{name} ({id_})"
    table.rows[row_idx].cells[1].text = id_
    table.rows[row_idx].cells[2].text = desc

doc.add_paragraph()

# Administrative Modules
p = doc.add_paragraph()
run = p.add_run("ADMINISTRATIVE MODULES")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'

for i, header in enumerate(headers):
    cell = table2.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

admin_modules = [
    ('Finance', 'M04', 'Billing - GST, TPA/Insurance'),
    ('HR & Payroll', 'M05', 'Attendance, Leave, Payroll'),
    ('Inventory', 'M06', 'Procurement, Vendor Mgmt')
]

for row_idx, (name, id_, desc) in enumerate(admin_modules, start=1):
    table2.rows[row_idx].cells[0].text = f"{name} ({id_})"
    table2.rows[row_idx].cells[1].text = id_
    table2.rows[row_idx].cells[2].text = desc

doc.add_paragraph()

# Platform Modules
p = doc.add_paragraph()
run = p.add_run("PLATFORM MODULES")
run.font.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0, 51, 102)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Table Grid'

for i, header in enumerate(headers):
    cell = table3.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

platform_modules = [
    ('BI & Reports', 'M07', 'Dashboards, ETL, Power BI'),
    ('Patient Portal', 'M08', 'PWA - Appointments, Reports'),
    ('Audit & Security', 'M09', 'Logs, DISHA compliance'),
    ('Appointments', 'M11', 'Scheduling, Reminders')
]

for row_idx, (name, id_, desc) in enumerate(platform_modules, start=1):
    table3.rows[row_idx].cells[0].text = f"{name} ({id_})"
    table3.rows[row_idx].cells[1].text = id_
    table3.rows[row_idx].cells[2].text = desc

# 3. Technology Stack
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("3. Technology Stack")
run.font.size = Pt(16)
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("Frontend: ")
run.font.bold = True
p.add_run("React 18 + TypeScript, Ant Design Pro, Vite, PWA")

p = doc.add_paragraph()
run = p.add_run("Backend: ")
run.font.bold = True
p.add_run("Spring Boot 3.x (Java 21), DDD + CQRS, Spring Cloud Gateway")

p = doc.add_paragraph()
run = p.add_run("Database: ")
run.font.bold = True
p.add_run("PostgreSQL 15 (Multi-tenant), Redis 7, Elasticsearch 8, MinIO/S3")

p = doc.add_paragraph()
run = p.add_run("Infrastructure: ")
run.font.bold = True
p.add_run("Kubernetes (EKS), Terraform, GitHub Actions, Apache Kafka")

# 4. Database Design
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("4. Database Design")
run.font.size = Pt(16)
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("Multi-Tenancy: ")
run.font.bold = True
p.add_run("Schema-isolated (separate PostgreSQL schema per tenant)")

table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Table Grid'

dbs = [
    ('PostgreSQL 15', 'Primary Data Store - Schema isolation, Row-level security'),
    ('Redis 7', 'Caching & Pub/Sub - Sessions, Rate limiting'),
    ('Elasticsearch 8', 'Analytics - Event indexing, Full-text search'),
    ('MinIO/S3', 'Document Storage - Reports, Prescriptions')
]

for row_idx, (db, desc) in enumerate(dbs):
    table4.rows[row_idx].cells[0].text = db
    run = table4.rows[row_idx].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    table4.rows[row_idx].cells[1].text = desc

# 5. RBAC
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("5. User Roles & Permissions (6-Tier RBAC)")
run.font.size = Pt(16)
run.font.bold = True

table5 = doc.add_table(rows=7, cols=4)
table5.style = 'Table Grid'

headers5 = ['Level', 'Role', 'Scope', 'Access']
for i, header in enumerate(headers5):
    cell = table5.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

roles = [
    ('L0', 'Super Admin', 'Platform', 'Multi-tenant management'),
    ('L1', 'Tenant Admin', 'Organization', 'Full hospital/chain'),
    ('L2', 'Branch Admin', 'Branch', 'Single branch'),
    ('L3', 'Department Head', 'Department', 'Department oversight'),
    ('L4', 'Clinical Staff', 'Department', 'Doctor, Nurse, Pharmacist'),
    ('L5', 'Support Staff', 'Assigned', 'Front desk, Billing')
]

for row_idx, (level, role, scope, access) in enumerate(roles, start=1):
    table5.rows[row_idx].cells[0].text = level
    table5.rows[row_idx].cells[1].text = role
    table5.rows[row_idx].cells[2].text = scope
    table5.rows[row_idx].cells[3].text = access

# 6. Security
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("6. Security Requirements")
run.font.size = Pt(16)
run.font.bold = True

table6 = doc.add_table(rows=9, cols=2)
table6.style = 'Table Grid'

security_items = [
    ('Authentication', 'JWT + Refresh tokens, MFA mandatory for L1-L3'),
    ('Session Management', 'Single active session for sensitive roles'),
    ('Failed Login', '5 failures -> account lockout + Admin alert'),
    ('Audit Logging', 'Immutable append-only logs, 7-year retention'),
    ('Data Encryption', 'AES-256 for backups, data at rest'),
    ('Compliance', 'DISHA (India healthcare), OWASP hardened')
]

for row_idx, (feature, impl) in enumerate(security_items):
    table6.rows[row_idx].cells[0].text = feature
    run = table6.rows[row_idx].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    table6.rows[row_idx].cells[1].text = impl

# 7. Delivery Phases
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("7. Delivery Phases (13 Phases)")
run.font.size = Pt(16)
run.font.bold = True

table7 = doc.add_table(rows=14, cols=4)
table7.style = 'Table Grid'

headers7 = ['Phase', 'Name', 'Duration', 'Key Deliverables']
for i, header in enumerate(headers7):
    cell = table7.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

phases = [
    ('P0', 'Foundation', '-', 'Architecture, ADRs'),
    ('P1', 'Multi-Tenant Core', '6 weeks', 'Schema, JWT, MFA'),
    ('P2', 'RBAC & Branch', '4 weeks', 'Roles, Permissions'),
    ('P3', 'OPD Module', '5 weeks', 'Patient, EMR, E-Rx'),
    ('P4', 'Pharmacy', '5 weeks', 'FIFO, Drug interaction'),
    ('P5', 'Laboratory', '4 weeks', 'Sample, Analyzers'),
    ('P6', 'Finance', '5 weeks', 'Billing, GST'),
    ('P7', 'HR & Payroll', '5 weeks', 'Attendance, Payroll'),
    ('P8', 'Inventory', '4 weeks', 'Procurement'),
    ('P9', 'BI & Reports', '4 weeks', 'Dashboards'),
    ('P10', 'Audit & Security', '3 weeks', 'Compliance'),
    ('P11', 'E2E QA', '5 weeks', 'Testing'),
    ('P12', 'Go-Live', '3 weeks', 'Deployment')
]

for row_idx, (phase, name, duration, deliverables) in enumerate(phases):
    table7.rows[row_idx].cells[0].text = phase
    table7.rows[row_idx].cells[1].text = name
    table7.rows[row_idx].cells[2].text = duration
    table7.rows[row_idx].cells[3].text = deliverables

# 8. OPD Details
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("8. OPD Module (M01) - Detailed Features")
run.font.size = Pt(16)
run.font.bold = True

features = [
    'UHID (Universal Health ID) generation',
    'Patient registration with demographic data',
    'EMR consultation with SOAP format notes',
    'Vitals entry (BP, pulse, temp, SpO2, BMI)',
    'ICD-10 diagnosis picker',
    'Real-time allergy warnings',
    'E-prescription with drug interaction (Rxnorm)',
    'Auto-save every 30 seconds'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 9. Pharmacy Details
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("9. Pharmacy Module (M02)")
run.font.size = Pt(16)
run.font.bold = True

features = [
    'FIFO batch dispensing engine',
    'Full batch traceability',
    'Real-time drug interaction alerts',
    'Automated stock reorder triggers',
    'Kafka-driven billing integration',
    'Prescription validation workflow'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 10. Laboratory Details
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("10. Laboratory Module (M03)")
run.font.size = Pt(16)
run.font.bold = True

features = [
    'Barcode-tracked sample management',
    'HL7/ASTM analyzer integration',
    'Dual-validation pathology workflow',
    'PDF report with secure S3 URLs',
    'Critical value alerts',
    'Quality control charts'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 11. Finance Details
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("11. Finance Module (M04)")
run.font.size = Pt(16)
run.font.bold = True

features = [
    'Multi-module bill aggregation',
    'GST compliance',
    'TPA/Insurance pre-authorisation',
    'Payment gateway integration',
    'Day-end close automation',
    'Financial reports (P&L, Balance Sheet)'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 12. HR Details
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("12. HR & Payroll Module (M05)")
run.font.size = Pt(16)
run.font.bold = True

features = [
    'Biometric attendance (ZKTeco/Suprema)',
    'Shift scheduling',
    'Leave management (CL, SL, EL, Maternity)',
    'Payroll: Basic + HRA + TA + Allowances',
    'Deductions: PF (12%+12%), ESIC, PT, TDS'
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 13. Integrations
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("13. Integration Requirements")
run.font.size = Pt(16)
run.font.bold = True

table8 = doc.add_table(rows=9, cols=3)
table8.style = 'Table Grid'

headers8 = ['Integration', 'Technology', 'Purpose']
for i, header in enumerate(headers8):
    cell = table8.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

integrations = [
    ('Biometric', 'ZKTeco/Suprema', 'Attendance'),
    ('Payment', 'Razorpay/Paytm', 'Payments'),
    ('SMS/WhatsApp', 'Twilio/MSG91', 'Notifications'),
    ('Email', 'SendGrid/SES', 'Communications'),
    ('BI Tools', 'Power BI/Tableau', 'Analytics'),
    ('Lab Analyzers', 'HL7/ASTM', 'Diagnostics'),
    ('Alerting', 'PagerDuty', 'Incidents'),
    ('Issue Tracking', 'Jira', 'Tickets')
]

for row_idx, (integ, tech, purpose) in enumerate(integrations, start=1):
    table8.rows[row_idx].cells[0].text = integ
    table8.rows[row_idx].cells[1].text = tech
    table8.rows[row_idx].cells[2].text = purpose

# 14. Kafka Events
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("14. Kafka Event Architecture")
run.font.size = Pt(16)
run.font.bold = True

table9 = doc.add_table(rows=9, cols=3)
table9.style = 'Table Grid'

headers9 = ['Event', 'Flow', 'Purpose']
for i, header in enumerate(headers9):
    cell = table9.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

events = [
    ('prescription.created', 'OPD -> Pharmacy', 'Process Rx'),
    ('laborder.created', 'OPD -> Lab', 'Create task'),
    ('drug.dispensed', 'Pharmacy -> Finance', 'Charge patient'),
    ('result.validated', 'Lab -> Finance', 'Update bill'),
    ('payroll.run', 'HR -> Finance', 'Process salary'),
    ('day-close.done', 'Finance', 'Trigger ETL'),
    ('bed.allocated', 'IPD -> Finance', 'Room charges'),
    ('appointment.booked', 'Appointments -> OPD', 'Check-in')
]

for row_idx, (event, flow, purpose) in enumerate(events, start=1):
    table9.rows[row_idx].cells[0].text = event
    table9.rows[row_idx].cells[1].text = flow
    table9.rows[row_idx].cells[2].text = purpose

# 15. Performance
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("15. Performance Requirements")
run.font.size = Pt(16)
run.font.bold = True

table10 = doc.add_table(rows=7, cols=2)
table10.style = 'Table Grid'

perf = [
    ('Uptime', '>99.9%'),
    ('API Response P95', '<300ms'),
    ('Dashboard Load', '<100ms'),
    ('Real-time Widgets', '<200ms'),
    ('Event Propagation', '<1 second'),
    ('Max Tenants', '50 concurrent')
]

for row_idx, (metric, target) in enumerate(perf):
    table10.rows[row_idx].cells[0].text = metric
    run = table10.rows[row_idx].cells[0].paragraphs[0].runs[0]
    run.font.bold = True
    table10.rows[row_idx].cells[1].text = target

# 16. Risks
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("16. Risk Mitigation")
run.font.size = Pt(16)
run.font.bold = True

table11 = doc.add_table(rows=7, cols=3)
table11.style = 'Table Grid'

headers11 = ['Risk', 'Severity', 'Mitigation']
for i, header in enumerate(headers11):
    cell = table11.rows[0].cells[i]
    cell.text = header
    run = cell.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    set_cell_shading(cell, '003366')

risks = [
    ('DISHA Compliance', 'HIGH', 'Checklist at design phase'),
    ('Multi-Tenant Perf', 'HIGH', 'Load testing, connection pools'),
    ('Dev Availability', 'HIGH', 'Documentation, cross-training'),
    ('Scope Creep', 'HIGH', 'SOW freeze, Change Control'),
    ('Kafka Complexity', 'MEDIUM', 'DLQ, idempotent consumers'),
    ('API Instability', 'MEDIUM', 'Circuit breakers, fallbacks')
]

for row_idx, (risk, severity, mitigation) in enumerate(risks, start=1):
    table11.rows[row_idx].cells[0].text = risk
    table11.rows[row_idx].cells[1].text = severity
    table11.rows[row_idx].cells[2].text = mitigation

# 17. Roadmap
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("17. Implementation Roadmap")
run.font.size = Pt(16)
run.font.bold = True

p = doc.add_paragraph()
run = p.add_run("PHASE 1: Foundation (Weeks 1-6)")
run.font.bold = True
run.font.size = Pt(13)

phase1 = [
    'Set up Kubernetes on EKS',
    'Configure PostgreSQL schema isolation',
    'Implement JWT + MFA authentication',
    'Set up Kafka cluster',
    'Configure Redis and Elasticsearch'
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("PHASE 2: Core Modules (Weeks 7-26)")
run.font.bold = True
run.font.size = Pt(13)

phase2 = [
    'Implement RBAC system',
    'Build OPD module',
    'Build Pharmacy module',
    'Build Laboratory module',
    'Build Finance module',
    'Build HR & Payroll'
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("PHASE 3: Platform (Weeks 27-38)")
run.font.bold = True
run.font.size = Pt(13)

phase3 = [
    'Build BI & Reports',
    'Implement Patient Portal PWA',
    'Build IPD module',
    'Implement Audit & Security',
    'Set up Kafka event flow'
]
for item in phase3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("PHASE 4: Launch (Weeks 39-46)")
run.font.bold = True
run.font.size = Pt(13)

phase4 = [
    'E2E testing with Playwright',
    'Load testing with k6',
    'Security penetration testing',
    'DISHA compliance audit',
    'Production deployment',
    'Pilot onboarding + 90-day hypercare'
]
for item in phase4:
    doc.add_paragraph(item, style='List Bullet')

# 18. Next Steps
doc.add_page_break()
heading = doc.add_paragraph()
run = heading.add_run("18. Next Steps")
run.font.size = Pt(16)
run.font.bold = True

steps = [
    'Finalize team composition',
    'Set up development environment',
    'Create technical design documents',
    'Define API contracts',
    'Set up CI/CD pipelines',
    'Begin Phase 1 development',
    'Schedule weekly sprint reviews'
]

for step in steps:
    doc.add_paragraph(step, style='List Bullet')

# Footer
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("--- End of HMS Detailed Build Plan ---")
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

doc.save(OUTPUT_PATH)
print(f"Successfully generated: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH) / 1024:.2f} KB")